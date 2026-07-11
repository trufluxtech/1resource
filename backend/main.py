import csv
import hashlib
import hmac
import io
import json
import os
import re
import secrets
import time
import sqlite3
import smtplib
import socket
import uuid
import base64
import urllib.error
import urllib.request
from email.message import EmailMessage
from collections import Counter, defaultdict
from datetime import datetime, timedelta
from typing import Any, Dict, List, Optional, Tuple
from urllib.parse import quote_plus
from xml.sax.saxutils import escape

from fastapi import Depends, FastAPI, File, Form, Header, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.lib.utils import ImageReader
from reportlab.platypus import Image as RLImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

APP_NAME = "1Resource"
APP_VERSION = "Production 1.13"
APP_ENV = os.environ.get("APP_ENV", "development").strip().lower()
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.environ.get("DATA_DIR", os.path.join(BASE_DIR, "data"))
UPLOAD_DIR = os.path.join(DATA_DIR, "uploads")
DB_PATH = os.environ.get("DATABASE_PATH", os.path.join(DATA_DIR, "resume_bank.db"))
DATABASE_URL = os.environ.get("DATABASE_URL", "").strip()
USE_POSTGRES = DATABASE_URL.startswith(("postgres://", "postgresql://"))
if APP_ENV in {"production", "prod"} and not USE_POSTGRES:
    raise RuntimeError("Production mode requires PostgreSQL. Set Railway DATABASE_URL=${{Postgres.DATABASE_URL}}. SQLite is allowed only for local development.")
SESSION_TTL_HOURS = int(os.environ.get("SESSION_TTL_HOURS", "8"))
MAX_UPLOAD_MB = int(os.environ.get("MAX_UPLOAD_MB", "10"))
MAX_UPLOAD_BYTES = MAX_UPLOAD_MB * 1024 * 1024
ALLOWED_UPLOAD_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
ALLOWED_PHOTO_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp"}
MAX_PHOTO_MB = int(os.environ.get("MAX_PHOTO_MB", "5"))
MAX_PHOTO_BYTES = MAX_PHOTO_MB * 1024 * 1024
ALLOWED_ORIGINS = [
    origin.strip() for origin in os.environ.get(
        "ALLOWED_ORIGINS",
        "http://localhost:5173,http://127.0.0.1:5173,http://localhost:8000,http://127.0.0.1:8000"
    ).split(",") if origin.strip()
]
PUBLIC_TOKEN_PATTERN = re.compile(r"^[A-Za-z0-9_\-]{32,128}$")
SMTP_HOST = os.environ.get("SMTP_HOST", "").strip()
SMTP_PORT = int(os.environ.get("SMTP_PORT", "587"))
SMTP_USERNAME = os.environ.get("SMTP_USERNAME", "").strip()
SMTP_PASSWORD = os.environ.get("SMTP_PASSWORD", "").strip()
SMTP_FROM_EMAIL = os.environ.get("SMTP_FROM_EMAIL", SMTP_USERNAME).strip()
SMTP_FROM_NAME = os.environ.get("SMTP_FROM_NAME", "1Resource Team").strip()
SMTP_USE_TLS = os.environ.get("SMTP_USE_TLS", "true").lower() not in {"0", "false", "no"}
EMAIL_PROVIDER = os.environ.get("EMAIL_PROVIDER", "smtp").strip().lower()
MAILJET_API_KEY = os.environ.get("MAILJET_API_KEY", os.environ.get("MJ_APIKEY_PUBLIC", "")).strip()
MAILJET_SECRET_KEY = os.environ.get("MAILJET_SECRET_KEY", os.environ.get("MJ_APIKEY_PRIVATE", "")).strip()
EMAIL_FROM = os.environ.get("EMAIL_FROM", SMTP_FROM_EMAIL).strip()
EMAIL_FROM_NAME = os.environ.get("EMAIL_FROM_NAME", SMTP_FROM_NAME).strip()
MAILJET_SEND_URL = os.environ.get("MAILJET_SEND_URL", "https://api.mailjet.com/v3.1/send").strip()
RATE_BUCKETS: Dict[str, List[float]] = {}

os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(UPLOAD_DIR, exist_ok=True)

app = FastAPI(title=APP_NAME, version=APP_VERSION)
app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_origin_regex=r"https?://(localhost|127\.0\.0\.1|0\.0\.0\.0|192\.168\.[0-9]+\.[0-9]+)(:[0-9]+)?",
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "PATCH", "DELETE", "OPTIONS"],
    allow_headers=["Authorization", "Content-Type", "X-Requested-With"],
)


def get_client_ip(request: Request) -> str:
    forwarded = request.headers.get("x-forwarded-for", "").split(",")[0].strip()
    return forwarded or (request.client.host if request.client else "unknown")


def check_rate_limit(key: str, limit: int, window_seconds: int = 60) -> None:
    now = time.time()
    bucket = [t for t in RATE_BUCKETS.get(key, []) if now - t < window_seconds]
    if len(bucket) >= limit:
        raise HTTPException(status_code=429, detail="Too many requests. Please try again shortly.")
    bucket.append(now)
    RATE_BUCKETS[key] = bucket


@app.middleware("http")
async def security_middleware(request: Request, call_next):
    path = request.url.path
    ip = get_client_ip(request)
    if path.startswith("/api/login"):
        check_rate_limit(f"login:{ip}", 15, 60)
    elif path.startswith("/api/public-upload"):
        check_rate_limit(f"public-upload:{ip}", 30, 60)
    elif path.startswith("/api/"):
        check_rate_limit(f"api:{ip}", 300, 60)
    response = await call_next(request)
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["X-Frame-Options"] = "DENY"
    response.headers["Referrer-Policy"] = "no-referrer"
    response.headers["Permissions-Policy"] = "camera=(), microphone=(), geolocation=()"
    response.headers["Cache-Control"] = "no-store" if path.startswith("/api/") else "no-cache"
    if request.url.scheme == "https":
        response.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"
    return response

SKILL_ALIASES: Dict[str, List[str]] = {
    "React": ["react", "reactjs", "react.js", "redux", "next.js", "nextjs", "frontend", "front end"],
    "Node.js": ["node", "node.js", "nodejs", "express", "express.js", "nestjs", "nest js", "backend", "back end"],
    "Python": ["python", "pandas", "numpy"],
    "FastAPI": ["fastapi", "fast api"],
    "Django": ["django", "drf", "django rest framework"],
    "Java": ["java", "spring boot", "spring", "springboot", "microservices"],
    "JavaScript": ["javascript", "js", "es6", "ecmascript"],
    "TypeScript": ["typescript", "ts"],
    "SQL": ["sql", "mysql", "postgres", "postgresql", "sqlite", "oracle", "sql server", "mssql", "plsql", "pl/sql"],
    "MongoDB": ["mongodb", "mongo", "nosql"],
    "AWS": ["aws", "amazon web services", "lambda", "ec2", "s3", "cloudwatch"],
    "Azure": ["azure", "azure devops", "azure functions", "aks"],
    "Salesforce": ["salesforce", "apex", "lwc", "lightning web component", "lightning", "flow", "flows", "soql", "salesforce health cloud", "health cloud", "salesforce shield"],
    "HL7/FHIR": ["hl7", "fhir", "hl7 fhir", "healthcare interoperability"],
    "Docker": ["docker", "container", "containers", "containerization"],
    "Kubernetes": ["kubernetes", "k8s", "helm"],
    "Terraform": ["terraform", "iac", "infrastructure as code"],
    "DevOps": ["devops", "ci/cd", "cicd", "github actions", "jenkins", "gitlab ci", "azure pipelines", "release pipeline"],
    "QA Automation": ["qa automation", "selenium", "playwright", "cypress", "test automation", "automation testing"],
    "Manual Testing": ["manual testing", "test cases", "uat", "functional testing"],
    "UI/UX": ["ui/ux", "ux", "figma", "wireframe", "prototype", "design system"],
    "Power BI": ["power bi", "powerbi", "dax"],
    "Tableau": ["tableau"],
    "Data Engineering": ["data engineering", "data engineer", "etl", "pipeline", "airflow", "spark", "databricks", "data pipeline", "data warehouse", "bigquery", "snowflake"],
    "AI/ML": ["machine learning", "ml", "ai", "llm", "rag", "genai", "generative ai", "model", "deep learning"],
    "Business Analyst": ["business analyst", "business analysis", "user stories", "brd", "frd", "requirements", "stakeholder"],
    "Project Management": ["project manager", "project management", "scrum", "agile", "pmp", "delivery manager", "program manager"],
    "Mobile": ["react native", "flutter", "android", "ios", "mobile", "kotlin", "swift"],
}

SKILL_CLUSTERS: Dict[str, List[str]] = {
    "Frontend": ["React", "JavaScript", "TypeScript", "UI/UX", "Mobile"],
    "Backend": ["Node.js", "Python", "FastAPI", "Django", "Java", "SQL", "MongoDB", "Salesforce", "HL7/FHIR"],
    "Data & AI": ["AI/ML", "Data Engineering", "Power BI", "Tableau", "Python", "SQL"],
    "Cloud & DevOps": ["AWS", "Azure", "Docker", "Kubernetes", "Terraform", "DevOps"],
    "Quality": ["QA Automation", "Manual Testing"],
    "Consulting & Delivery": ["Business Analyst", "Project Management"],
}

SKILL_TO_CLUSTER: Dict[str, str] = {
    skill: cluster for cluster, skills in SKILL_CLUSTERS.items() for skill in skills
}


def skill_cluster(skill: str) -> str:
    return SKILL_TO_CLUSTER.get(skill, skill)


DOMAIN_TERMS = [
    "Banking", "Insurance", "Government", "GovTech", "Retail", "SaaS", "Healthcare", "Education",
    "Manufacturing", "Logistics", "Travel", "Telecom", "Media", "ERP", "CRM", "FinTech", "Railways"
]


def now_iso() -> str:
    return datetime.utcnow().replace(microsecond=0).isoformat() + "Z"


def today_date() -> str:
    return datetime.utcnow().date().isoformat()


def parse_iso(value: str) -> datetime:
    return datetime.fromisoformat(value.replace("Z", ""))


def parse_optional_iso(value: Any) -> Optional[datetime]:
    try:
        text = str(value or "").strip()
        if not text:
            return None
        return parse_iso(text)
    except Exception:
        return None


def assert_public_token(token: str) -> None:
    if not PUBLIC_TOKEN_PATTERN.match(token or ""):
        raise HTTPException(status_code=400, detail="Invalid upload link token")


class PgCursor:
    def __init__(self, cursor, lastrowid: Optional[int] = None):
        self.cursor = cursor
        self.lastrowid = lastrowid

    def fetchone(self):
        return self.cursor.fetchone()

    def fetchall(self):
        return self.cursor.fetchall()


class PgConnection:
    """Small compatibility wrapper so existing SQLite-style code can run on PostgreSQL."""
    is_postgres = True

    def __init__(self):
        try:
            import psycopg2
            import psycopg2.extras
        except ImportError as exc:
            raise RuntimeError("PostgreSQL mode requires psycopg2-binary. Run: pip install -r backend/requirements.txt") from exc
        self.psycopg2 = psycopg2
        self.conn = psycopg2.connect(DATABASE_URL, cursor_factory=psycopg2.extras.RealDictCursor)
        self.conn.autocommit = False

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc, tb):
        try:
            if exc_type:
                self.conn.rollback()
            else:
                self.conn.commit()
        finally:
            self.conn.close()

    def _convert_sql(self, sql: str, add_returning: bool = True) -> Tuple[str, bool]:
        s = sql.strip()
        while s.endswith(";"):
            s = s[:-1].strip()

        s = re.sub(r"INTEGER\s+PRIMARY\s+KEY\s+AUTOINCREMENT", "SERIAL PRIMARY KEY", s, flags=re.I)

        # SQLite syntax compatibility.
        if re.match(r"INSERT\s+OR\s+IGNORE\s+INTO\s+company_profile", s, flags=re.I):
            s = re.sub(r"INSERT\s+OR\s+IGNORE\s+INTO", "INSERT INTO", s, count=1, flags=re.I)
            s += " ON CONFLICT (id) DO NOTHING"
        elif re.match(r"INSERT\s+OR\s+IGNORE\s+INTO", s, flags=re.I):
            s = re.sub(r"INSERT\s+OR\s+IGNORE\s+INTO", "INSERT INTO", s, count=1, flags=re.I)
            s += " ON CONFLICT DO NOTHING"

        needs_returning = False
        if add_returning:
            first_words = s.lstrip().upper()
            no_id_tables = ("SESSIONS",)
            m = re.match(r"INSERT\s+INTO\s+([A-Za-z_][A-Za-z0-9_]*)", s, flags=re.I)
            if m and "RETURNING" not in first_words:
                table = m.group(1).upper()
                if table not in no_id_tables:
                    needs_returning = True
                    s += " RETURNING id"

        # Convert SQLite placeholders to psycopg2 placeholders.
        s = s.replace("?", "%s")
        return s, needs_returning

    def execute(self, sql: str, params: Any = None):
        params = [] if params is None else params
        sql2, wants_rowid = self._convert_sql(sql, add_returning=True)
        cur = self.conn.cursor()
        try:
            cur.execute(sql2, params)
            lastrowid = None
            if wants_rowid:
                returned = cur.fetchone()
                if returned and "id" in returned:
                    lastrowid = returned["id"]
            return PgCursor(cur, lastrowid)
        except Exception as exc:
            self.conn.rollback()
            if exc.__class__.__name__ in {"UniqueViolation", "IntegrityError"}:
                raise sqlite3.IntegrityError(str(exc))
            raise

    def executescript(self, script: str):
        cur = self.conn.cursor()
        statements = [stmt.strip() for stmt in script.split(";") if stmt.strip()]
        for stmt in statements:
            sql2, _ = self._convert_sql(stmt, add_returning=False)
            cur.execute(sql2)
        return PgCursor(cur)

    def commit(self):
        self.conn.commit()

    def rollback(self):
        self.conn.rollback()

    def close(self):
        self.conn.close()


def get_db():
    if USE_POSTGRES:
        return PgConnection()
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def dict_row(row) -> Dict[str, Any]:
    if row is None:
        return {}
    return {k: row[k] for k in row.keys()}



def hash_password(password: str, salt: Optional[str] = None) -> str:
    salt = salt or secrets.token_hex(16)
    digest = hashlib.pbkdf2_hmac("sha256", password.encode(), salt.encode(), 150000).hex()
    return f"{salt}${digest}"


def verify_password(password: str, stored: str) -> bool:
    try:
        salt, digest = stored.split("$", 1)
        calculated = hash_password(password, salt).split("$", 1)[1]
        return hmac.compare_digest(calculated, digest)
    except Exception:
        return False


def enforce_password_policy(password: str) -> None:
    if len(password or "") < 8:
        raise HTTPException(status_code=400, detail="Password must be at least 8 characters long")
    if not re.search(r"[A-Za-z]", password) or not re.search(r"\d", password):
        raise HTTPException(status_code=400, detail="Password must contain at least one letter and one number")
    if password in {"admin123", "password", "Welcome123"}:
        raise HTTPException(status_code=400, detail="Please use a stronger password")


def session_expiry() -> str:
    return (datetime.utcnow() + timedelta(hours=SESSION_TTL_HOURS)).replace(microsecond=0).isoformat() + "Z"


def table_columns(conn, table: str) -> List[str]:
    try:
        if getattr(conn, "is_postgres", False):
            rows = conn.execute(
                "SELECT column_name AS name FROM information_schema.columns WHERE table_schema='public' AND table_name=%s",
                (table,),
            ).fetchall()
            return [r["name"] for r in rows]
        return [r["name"] for r in conn.execute(f"PRAGMA table_info({table})").fetchall()]
    except Exception:
        return []


def add_column_if_missing(conn, table: str, column: str, definition: str) -> None:
    if column not in table_columns(conn, table):
        conn.execute(f"ALTER TABLE {table} ADD COLUMN {column} {definition}")


def init_db() -> None:
    with get_db() as conn:
        conn.executescript(
            """
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE NOT NULL,
                full_name TEXT NOT NULL,
                email TEXT,
                phone TEXT,
                role TEXT NOT NULL DEFAULT 'Recruiter',
                password_hash TEXT NOT NULL,
                is_active INTEGER NOT NULL DEFAULT 1,
                failed_attempts INTEGER NOT NULL DEFAULT 0,
                locked_until TEXT,
                force_password_change INTEGER NOT NULL DEFAULT 0,
                password_changed_at TEXT,
                last_login_at TEXT,
                created_at TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS sessions (
                token TEXT PRIMARY KEY,
                user_id INTEGER NOT NULL,
                created_at TEXT NOT NULL,
                expires_at TEXT,
                last_seen_at TEXT,
                user_agent TEXT,
                ip_address TEXT,
                FOREIGN KEY(user_id) REFERENCES users(id)
            );

            CREATE TABLE IF NOT EXISTS candidates (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                candidate_code TEXT UNIQUE NOT NULL,
                full_name TEXT NOT NULL,
                email TEXT,
                phone TEXT,
                location TEXT,
                current_status TEXT DEFAULT 'Available',
                availability_date TEXT,
                available_by_date TEXT,
                notice_period_days INTEGER DEFAULT 0,
                employment_type TEXT DEFAULT 'Contract',
                source TEXT,
                recruiter_owner TEXT,
                total_experience REAL DEFAULT 0,
                relevant_experience REAL DEFAULT 0,
                primary_skill TEXT,
                secondary_skills TEXT,
                domain_exposure TEXT,
                proficiency TEXT DEFAULT 'Intermediate',
                certifications TEXT,
                portfolio_url TEXT,
                current_company TEXT,
                previous_companies TEXT,
                project_details TEXT,
                photo_file_name TEXT,
                photo_file_path TEXT,
                expected_rate REAL DEFAULT 0,
                negotiated_rate REAL DEFAULT 0,
                internal_level TEXT DEFAULT 'L2 - Mid-level',
                resume_text TEXT,
                resume_file_name TEXT,
                resume_file_path TEXT,
                status TEXT DEFAULT 'New',
                created_date TEXT,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS assessments (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                candidate_id INTEGER NOT NULL,
                technical_score INTEGER DEFAULT 0,
                project_score INTEGER DEFAULT 0,
                practical_score INTEGER DEFAULT 0,
                communication_score INTEGER DEFAULT 0,
                client_readiness_score INTEGER DEFAULT 0,
                cost_fitment_score INTEGER DEFAULT 0,
                availability_score INTEGER DEFAULT 0,
                total_score INTEGER DEFAULT 0,
                evaluator_name TEXT,
                recommendation TEXT,
                remarks TEXT,
                created_at TEXT NOT NULL,
                FOREIGN KEY(candidate_id) REFERENCES candidates(id) ON DELETE CASCADE
            );

            CREATE TABLE IF NOT EXISTS resumes (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                candidate_id INTEGER NOT NULL,
                role_title TEXT,
                role_definition TEXT,
                resume_file_name TEXT,
                resume_file_path TEXT,
                resume_text TEXT,
                extracted_json TEXT,
                fit_score INTEGER DEFAULT 0,
                rating_level TEXT,
                fake_risk_score INTEGER DEFAULT 0,
                fake_risk_level TEXT,
                fake_risk_reasons TEXT,
                skill_matches TEXT,
                skill_gaps TEXT,
                source TEXT,
                created_at TEXT NOT NULL,
                FOREIGN KEY(candidate_id) REFERENCES candidates(id) ON DELETE CASCADE
            );

            CREATE TABLE IF NOT EXISTS demand_requests (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                demand_code TEXT UNIQUE NOT NULL,
                client_name TEXT,
                project_name TEXT,
                role_title TEXT NOT NULL,
                role_definition TEXT,
                required_skills TEXT,
                domain TEXT,
                location TEXT,
                work_mode TEXT DEFAULT 'Hybrid',
                priority TEXT DEFAULT 'Medium',
                status TEXT DEFAULT 'Open',
                number_of_positions INTEGER DEFAULT 1,
                target_customer_rate REAL DEFAULT 0,
                max_internal_cost REAL DEFAULT 0,
                start_date TEXT,
                duration_weeks INTEGER DEFAULT 12,
                created_by TEXT,
                created_date TEXT,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS demand_shortlists (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                demand_id INTEGER NOT NULL,
                candidate_id INTEGER NOT NULL,
                match_score INTEGER DEFAULT 0,
                match_level TEXT,
                skill_matches TEXT,
                skill_gaps TEXT,
                commercial_fit TEXT,
                availability_fit TEXT,
                status TEXT DEFAULT 'Shortlisted',
                notes TEXT,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL,
                UNIQUE(demand_id, candidate_id),
                FOREIGN KEY(demand_id) REFERENCES demand_requests(id) ON DELETE CASCADE,
                FOREIGN KEY(candidate_id) REFERENCES candidates(id) ON DELETE CASCADE
            );

            CREATE TABLE IF NOT EXISTS demand_mcq_questions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                demand_id INTEGER NOT NULL,
                question_no INTEGER NOT NULL,
                skill TEXT,
                question_text TEXT NOT NULL,
                option_a TEXT NOT NULL,
                option_b TEXT NOT NULL,
                option_c TEXT NOT NULL,
                option_d TEXT NOT NULL,
                correct_option TEXT NOT NULL,
                explanation TEXT,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL,
                UNIQUE(demand_id, question_no),
                FOREIGN KEY(demand_id) REFERENCES demand_requests(id) ON DELETE CASCADE
            );

            CREATE TABLE IF NOT EXISTS public_upload_links (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                token TEXT UNIQUE NOT NULL,
                role_title TEXT,
                role_definition TEXT,
                candidate_id INTEGER,
                demand_id INTEGER,
                include_mcq INTEGER NOT NULL DEFAULT 1,
                created_by TEXT,
                expires_at TEXT NOT NULL,
                is_active INTEGER NOT NULL DEFAULT 1,
                revoked_at TEXT,
                revoked_by TEXT,
                created_at TEXT NOT NULL,
                used_at TEXT,
                FOREIGN KEY(candidate_id) REFERENCES candidates(id) ON DELETE SET NULL,
                FOREIGN KEY(demand_id) REFERENCES demand_requests(id) ON DELETE SET NULL
            );

            CREATE TABLE IF NOT EXISTS public_mcq_results (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                public_link_id INTEGER,
                demand_id INTEGER,
                candidate_id INTEGER,
                score REAL DEFAULT 0,
                total INTEGER DEFAULT 10,
                answered_count INTEGER DEFAULT 0,
                correct_count INTEGER DEFAULT 0,
                wrong_count INTEGER DEFAULT 0,
                negative_per_wrong REAL DEFAULT 0.25,
                percentage REAL DEFAULT 0,
                passed INTEGER DEFAULT 0,
                answers_json TEXT,
                created_at TEXT NOT NULL,
                FOREIGN KEY(public_link_id) REFERENCES public_upload_links(id) ON DELETE SET NULL,
                FOREIGN KEY(demand_id) REFERENCES demand_requests(id) ON DELETE SET NULL,
                FOREIGN KEY(candidate_id) REFERENCES candidates(id) ON DELETE SET NULL
            );

            CREATE TABLE IF NOT EXISTS company_profile (
                id INTEGER PRIMARY KEY CHECK (id = 1),
                company_name TEXT DEFAULT 'Truflux Technologies',
                company_number TEXT,
                tax_number TEXT,
                address TEXT,
                phone TEXT,
                email TEXT,
                website TEXT,
                logo_file_name TEXT,
                logo_file_path TEXT,
                updated_at TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS potential_clients (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                company_name TEXT NOT NULL,
                contact_name TEXT,
                email TEXT NOT NULL,
                phone TEXT,
                segment TEXT,
                status TEXT DEFAULT 'Prospect',
                notes TEXT,
                created_by TEXT,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS email_outreach_logs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                client_id INTEGER,
                user_id INTEGER,
                recipient_type TEXT NOT NULL,
                recipient_name TEXT,
                recipient_email TEXT NOT NULL,
                email_type TEXT NOT NULL,
                subject TEXT NOT NULL,
                body TEXT NOT NULL,
                status TEXT NOT NULL,
                error TEXT,
                created_by TEXT,
                created_at TEXT NOT NULL,
                sent_at TEXT,
                FOREIGN KEY(client_id) REFERENCES potential_clients(id) ON DELETE SET NULL,
                FOREIGN KEY(user_id) REFERENCES users(id) ON DELETE SET NULL
            );

            CREATE TABLE IF NOT EXISTS activity_logs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                actor TEXT,
                action TEXT NOT NULL,
                entity_type TEXT,
                entity_id INTEGER,
                details TEXT,
                created_at TEXT NOT NULL
            );
            """
        )
        # Migrations for users upgrading from previous ZIPs.
        add_column_if_missing(conn, "candidates", "ml_rating_score", "INTEGER DEFAULT 0")
        add_column_if_missing(conn, "candidates", "ml_rating_level", "TEXT")
        add_column_if_missing(conn, "candidates", "fake_risk_score", "INTEGER DEFAULT 0")
        add_column_if_missing(conn, "candidates", "fake_risk_level", "TEXT")
        add_column_if_missing(conn, "candidates", "fake_risk_reasons", "TEXT")
        add_column_if_missing(conn, "candidates", "skill_matches", "TEXT")
        add_column_if_missing(conn, "candidates", "skill_gaps", "TEXT")
        add_column_if_missing(conn, "candidates", "last_role_definition", "TEXT")
        add_column_if_missing(conn, "candidates", "resume_count", "INTEGER DEFAULT 0")
        add_column_if_missing(conn, "candidates", "available_by_date", "TEXT")
        add_column_if_missing(conn, "candidates", "notice_period_days", "INTEGER DEFAULT 0")
        add_column_if_missing(conn, "candidates", "created_date", "TEXT")
        add_column_if_missing(conn, "candidates", "current_company", "TEXT")
        add_column_if_missing(conn, "candidates", "previous_companies", "TEXT")
        add_column_if_missing(conn, "candidates", "project_details", "TEXT")
        add_column_if_missing(conn, "candidates", "photo_file_name", "TEXT")
        add_column_if_missing(conn, "candidates", "photo_file_path", "TEXT")
        add_column_if_missing(conn, "users", "email", "TEXT")
        add_column_if_missing(conn, "users", "phone", "TEXT")
        add_column_if_missing(conn, "users", "photo_file_name", "TEXT")
        add_column_if_missing(conn, "users", "photo_file_path", "TEXT")
        add_column_if_missing(conn, "users", "failed_attempts", "INTEGER NOT NULL DEFAULT 0")
        add_column_if_missing(conn, "users", "locked_until", "TEXT")
        add_column_if_missing(conn, "users", "force_password_change", "INTEGER NOT NULL DEFAULT 0")
        add_column_if_missing(conn, "users", "password_changed_at", "TEXT")
        add_column_if_missing(conn, "users", "last_login_at", "TEXT")
        add_column_if_missing(conn, "sessions", "expires_at", "TEXT")
        add_column_if_missing(conn, "sessions", "last_seen_at", "TEXT")
        add_column_if_missing(conn, "sessions", "user_agent", "TEXT")
        add_column_if_missing(conn, "sessions", "ip_address", "TEXT")
        add_column_if_missing(conn, "demand_shortlists", "availability_fit", "TEXT")
        add_column_if_missing(conn, "public_upload_links", "use_count", "INTEGER DEFAULT 0")
        add_column_if_missing(conn, "public_upload_links", "last_uploaded_at", "TEXT")
        add_column_if_missing(conn, "public_upload_links", "locked_note", "TEXT")
        add_column_if_missing(conn, "public_mcq_results", "correct_count", "INTEGER DEFAULT 0")
        add_column_if_missing(conn, "public_mcq_results", "wrong_count", "INTEGER DEFAULT 0")
        add_column_if_missing(conn, "public_mcq_results", "negative_per_wrong", "REAL DEFAULT 0.25")
        add_column_if_missing(conn, "public_mcq_results", "percentage", "REAL DEFAULT 0")
        add_column_if_missing(conn, "public_upload_links", "revoked_at", "TEXT")
        add_column_if_missing(conn, "public_upload_links", "revoked_by", "TEXT")
        add_column_if_missing(conn, "public_upload_links", "demand_id", "INTEGER")
        add_column_if_missing(conn, "public_upload_links", "include_mcq", "INTEGER NOT NULL DEFAULT 1")
        add_column_if_missing(conn, "company_profile", "company_name", "TEXT DEFAULT 'Truflux Technologies'")
        add_column_if_missing(conn, "company_profile", "company_number", "TEXT")
        add_column_if_missing(conn, "company_profile", "tax_number", "TEXT")
        add_column_if_missing(conn, "company_profile", "address", "TEXT")
        add_column_if_missing(conn, "company_profile", "phone", "TEXT")
        add_column_if_missing(conn, "company_profile", "email", "TEXT")
        add_column_if_missing(conn, "company_profile", "website", "TEXT")
        add_column_if_missing(conn, "company_profile", "logo_file_name", "TEXT")
        add_column_if_missing(conn, "company_profile", "logo_file_path", "TEXT")
        add_column_if_missing(conn, "potential_clients", "phone", "TEXT")
        add_column_if_missing(conn, "potential_clients", "segment", "TEXT")
        add_column_if_missing(conn, "potential_clients", "status", "TEXT DEFAULT 'Prospect'")
        add_column_if_missing(conn, "potential_clients", "notes", "TEXT")
        add_column_if_missing(conn, "potential_clients", "created_by", "TEXT")
        add_column_if_missing(conn, "email_outreach_logs", "user_id", "INTEGER")
        add_column_if_missing(conn, "email_outreach_logs", "recipient_type", "TEXT DEFAULT 'client'")
        add_column_if_missing(conn, "company_profile", "updated_at", "TEXT")
        add_column_if_missing(conn, "demand_requests", "created_date", "TEXT")
        # Backfill explicit created_date for old records so trend analytics uses a stable record-created date.
        conn.execute("UPDATE candidates SET created_date = substr(created_at, 1, 10) WHERE created_date IS NULL OR created_date = ''")
        conn.execute("UPDATE demand_requests SET created_date = substr(created_at, 1, 10) WHERE created_date IS NULL OR created_date = ''")
        conn.execute(
            """
            INSERT OR IGNORE INTO company_profile(id, company_name, company_number, tax_number, address, phone, email, website, updated_at)
            VALUES(1, 'Truflux Technologies', '', '', '', '', '', '', ?)
            """,
            (now_iso(),),
        )
        user_count = conn.execute("SELECT COUNT(*) AS c FROM users").fetchone()["c"]
        if user_count == 0:
            conn.execute(
                "INSERT INTO users(username, full_name, email, phone, role, password_hash, is_active, force_password_change, password_changed_at, created_at) VALUES(?,?,?,?,?,?,?,?,?,?)",
                ("Admin", "Truflux Admin", "", "", "Admin", hash_password("admin123"), 1, 0, now_iso(), now_iso()),
            )
        # Ensure a default recruiter profile exists for demos and resume-PDF authorized contact testing.
        conn.execute(
            """
            INSERT OR IGNORE INTO users(username, full_name, email, phone, role, password_hash, is_active, force_password_change, password_changed_at, created_at)
            VALUES(?,?,?,?,?,?,?,?,?,?)
            """,
            ("Recruiter", "Default Recruiter", "recruiter@truflux.ai", "+91 00000 00000", "Recruiter", hash_password("recruiter123"), 1, 0, now_iso(), now_iso()),
        )
        conn.commit()


def log_action(actor: str, action: str, entity_type: Optional[str] = None, entity_id: Optional[int] = None, details: Optional[str] = None) -> None:
    with get_db() as conn:
        conn.execute(
            "INSERT INTO activity_logs(actor, action, entity_type, entity_id, details, created_at) VALUES(?,?,?,?,?,?)",
            (actor, action, entity_type, entity_id, details, now_iso()),
        )
        conn.commit()


@app.on_event("startup")
def startup() -> None:
    init_db()


class LoginRequest(BaseModel):
    username: str
    password: str


class UserCreate(BaseModel):
    username: str
    full_name: str
    email: Optional[str] = ""
    phone: Optional[str] = ""
    role: str = "Recruiter"
    password: str = "Welcome1234"


class UserUpdate(BaseModel):
    username: str
    full_name: str
    email: Optional[str] = ""
    phone: Optional[str] = ""
    role: str = "Recruiter"


class UserProfileUpdate(BaseModel):
    full_name: Optional[str] = ""
    email: Optional[str] = ""
    phone: Optional[str] = ""


class CandidateBase(BaseModel):
    full_name: str
    email: Optional[str] = ""
    phone: Optional[str] = ""
    location: Optional[str] = ""
    current_status: Optional[str] = "Available"
    availability_date: Optional[str] = ""
    available_by_date: Optional[str] = ""
    notice_period_days: Optional[int] = 0
    employment_type: Optional[str] = "Contract"
    source: Optional[str] = ""
    recruiter_owner: Optional[str] = ""
    total_experience: Optional[float] = 0
    relevant_experience: Optional[float] = 0
    primary_skill: Optional[str] = ""
    secondary_skills: Optional[str] = ""
    domain_exposure: Optional[str] = ""
    proficiency: Optional[str] = "Intermediate"
    certifications: Optional[str] = ""
    portfolio_url: Optional[str] = ""
    current_company: Optional[str] = ""
    previous_companies: Optional[str] = ""
    project_details: Optional[str] = ""
    expected_rate: Optional[float] = 0
    negotiated_rate: Optional[float] = 0
    internal_level: Optional[str] = "L2 - Mid-level"
    resume_text: Optional[str] = ""
    status: Optional[str] = "New"


class CandidateCreate(CandidateBase):
    pass


class CandidateUpdate(CandidateBase):
    pass


class AssessmentCreate(BaseModel):
    technical_score: int = Field(0, ge=0, le=30)
    project_score: int = Field(0, ge=0, le=15)
    practical_score: int = Field(0, ge=0, le=20)
    communication_score: int = Field(0, ge=0, le=10)
    client_readiness_score: int = Field(0, ge=0, le=10)
    cost_fitment_score: int = Field(0, ge=0, le=10)
    availability_score: int = Field(0, ge=0, le=5)
    evaluator_name: Optional[str] = ""
    recommendation: Optional[str] = "Hold"
    remarks: Optional[str] = ""


class LinkCreate(BaseModel):
    role_title: str = ""
    role_definition: str = ""
    candidate_id: Optional[int] = None
    demand_id: Optional[int] = None
    include_mcq: bool = True


class PublicMcqSubmit(BaseModel):
    answers: Dict[str, str] = Field(default_factory=dict)


class CompanyProfileUpdate(BaseModel):
    company_name: str = "Truflux Technologies"
    company_number: Optional[str] = ""
    tax_number: Optional[str] = ""
    address: Optional[str] = ""
    phone: Optional[str] = ""
    email: Optional[str] = ""
    website: Optional[str] = ""


class PotentialClientBase(BaseModel):
    company_name: str
    contact_name: Optional[str] = ""
    email: str
    phone: Optional[str] = ""
    segment: Optional[str] = ""
    status: Optional[str] = "Prospect"
    notes: Optional[str] = ""


class PotentialClientCreate(PotentialClientBase):
    pass


class PotentialClientUpdate(PotentialClientBase):
    pass


class OutreachEmailRequest(BaseModel):
    client_ids: List[int] = Field(default_factory=list)
    user_ids: List[int] = Field(default_factory=list)
    email_type: str = "encourage"
    subject: str
    body: str


class DemandBase(BaseModel):
    client_name: Optional[str] = ""
    project_name: Optional[str] = ""
    role_title: str
    role_definition: Optional[str] = ""
    required_skills: Optional[str] = ""
    domain: Optional[str] = ""
    location: Optional[str] = ""
    work_mode: Optional[str] = "Hybrid"
    priority: Optional[str] = "Medium"
    status: Optional[str] = "Open"
    number_of_positions: Optional[int] = 1
    target_customer_rate: Optional[float] = 0
    max_internal_cost: Optional[float] = 0
    start_date: Optional[str] = ""
    duration_weeks: Optional[int] = 12


class DemandCreate(DemandBase):
    pass


class DemandUpdate(DemandBase):
    pass


class ShortlistUpdate(BaseModel):
    status: Optional[str] = "Shortlisted"
    notes: Optional[str] = ""



def validate_outreach_email(value: str) -> str:
    email = (value or "").strip()
    if not re.match(r"^[^@\s]+@[^@\s]+\.[^@\s]+$", email):
        raise HTTPException(status_code=400, detail=f"Invalid email address: {email or 'blank'}")
    return email


def validate_email_type(value: str) -> str:
    email_type = (value or "").strip().lower()
    allowed = {"encourage", "release_notes", "shutdown"}
    if email_type not in allowed:
        raise HTTPException(status_code=400, detail="Email type must be encourage, release_notes, or shutdown")
    return email_type


def smtp_is_configured() -> bool:
    return bool(SMTP_HOST and SMTP_FROM_EMAIL)


def friendly_smtp_error(exc: Exception) -> str:
    raw = str(exc)
    if isinstance(exc, socket.gaierror) or "Temporary failure in name resolution" in raw or "Name or service not known" in raw:
        return f"DNS could not resolve SMTP_HOST='{SMTP_HOST}'. Check Railway Variables and use your provider's exact SMTP server name, for example smtp.gmail.com, smtp.office365.com, or your email provider SMTP host."
    if isinstance(exc, TimeoutError) or "timed out" in raw.lower():
        return f"Connection to SMTP_HOST='{SMTP_HOST}' on port {SMTP_PORT} timed out. Check host, port and SMTP provider firewall/settings."
    if "Authentication" in raw or "Username and Password not accepted" in raw:
        return "SMTP authentication failed. Check SMTP_USERNAME and SMTP_PASSWORD. For Gmail/Workspace, use an App Password or approved SMTP relay."
    return raw[:500]


def email_provider_status() -> Dict[str, Any]:
    provider = EMAIL_PROVIDER or "smtp"
    if provider in {"mailjet", "mailjet_api"}:
        missing = []
        if not MAILJET_API_KEY:
            missing.append("MAILJET_API_KEY")
        if not MAILJET_SECRET_KEY:
            missing.append("MAILJET_SECRET_KEY")
        if not EMAIL_FROM:
            missing.append("EMAIL_FROM")
        return {
            "provider": "mailjet_api",
            "configured": not missing,
            "api_url": MAILJET_SEND_URL,
            "from_email": EMAIL_FROM or "",
            "from_name": EMAIL_FROM_NAME or "",
            "has_api_key": bool(MAILJET_API_KEY),
            "has_secret_key": bool(MAILJET_SECRET_KEY),
            "missing": missing,
        }
    return smtp_config_status()


def smtp_config_status() -> Dict[str, Any]:
    missing = []
    if not SMTP_HOST:
        missing.append("SMTP_HOST")
    if not SMTP_FROM_EMAIL:
        missing.append("SMTP_FROM_EMAIL")
    return {
        "provider": "smtp",
        "configured": bool(SMTP_HOST and SMTP_FROM_EMAIL),
        "host": SMTP_HOST or "",
        "port": SMTP_PORT,
        "from_email": SMTP_FROM_EMAIL or "",
        "from_name": SMTP_FROM_NAME or "",
        "tls": SMTP_USE_TLS,
        "has_username": bool(SMTP_USERNAME),
        "has_password": bool(SMTP_PASSWORD),
        "missing": missing,
    }


def test_smtp_connection() -> Dict[str, Any]:
    status = smtp_config_status()
    if not status["configured"]:
        return {"ok": False, **status, "message": f"SMTP is not configured. Missing: {', '.join(status['missing'])}"}
    try:
        socket.getaddrinfo(SMTP_HOST, SMTP_PORT)
        with smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=15) as smtp:
            if SMTP_USE_TLS:
                smtp.starttls()
            if SMTP_USERNAME and SMTP_PASSWORD:
                smtp.login(SMTP_USERNAME, SMTP_PASSWORD)
        return {"ok": True, **status, "message": "SMTP connection test passed."}
    except Exception as exc:
        return {"ok": False, **status, "message": friendly_smtp_error(exc)}


def test_mailjet_api_connection() -> Dict[str, Any]:
    status = email_provider_status()
    if not status["configured"]:
        return {"ok": False, **status, "message": f"Mailjet API is not configured. Missing: {', '.join(status['missing'])}"}
    # Safe connectivity/auth test: hit the Send API with an empty Messages array and read the HTTP response.
    payload = json.dumps({"Messages": []}).encode("utf-8")
    auth_raw = f"{MAILJET_API_KEY}:{MAILJET_SECRET_KEY}".encode("utf-8")
    req = urllib.request.Request(
        MAILJET_SEND_URL,
        data=payload,
        method="POST",
        headers={
            "Content-Type": "application/json",
            "Authorization": "Basic " + base64.b64encode(auth_raw).decode("ascii"),
        },
    )
    try:
        with urllib.request.urlopen(req, timeout=20) as resp:
            _ = resp.read()
            return {"ok": True, **status, "message": "Mailjet API HTTPS connection is reachable."}
    except urllib.error.HTTPError as exc:
        body = exc.read().decode("utf-8", errors="replace")[:500]
        # 400/422 means HTTPS + auth path is reachable; request is intentionally empty.
        if exc.code in {400, 422}:
            return {"ok": True, **status, "message": "Mailjet API HTTPS connection is reachable. The test request was rejected as expected because it did not contain a real message."}
        if exc.code in {401, 403}:
            return {"ok": False, **status, "message": "Mailjet API reached, but authentication failed. Check MAILJET_API_KEY and MAILJET_SECRET_KEY."}
        return {"ok": False, **status, "message": f"Mailjet API returned HTTP {exc.code}: {body}"}
    except Exception as exc:
        return {"ok": False, **status, "message": f"Mailjet API HTTPS connection failed: {str(exc)[:500]}"}


def test_email_connection() -> Dict[str, Any]:
    provider = EMAIL_PROVIDER or "smtp"
    if provider in {"mailjet", "mailjet_api"}:
        return test_mailjet_api_connection()
    return test_smtp_connection()


def send_mailjet_api_email(recipient_email: str, recipient_name: str, subject: str, body: str) -> Tuple[str, str]:
    status = email_provider_status()
    if not status["configured"]:
        return "queued_mailjet_api_not_configured", f"Mailjet API not configured. Missing: {', '.join(status['missing'])}"

    payload = {
        "Messages": [{
            "From": {"Email": EMAIL_FROM, "Name": EMAIL_FROM_NAME or "1Resource Team"},
            "To": [{"Email": recipient_email, "Name": recipient_name or recipient_email}],
            "Subject": subject,
            "TextPart": body,
        }]
    }
    auth_raw = f"{MAILJET_API_KEY}:{MAILJET_SECRET_KEY}".encode("utf-8")
    req = urllib.request.Request(
        MAILJET_SEND_URL,
        data=json.dumps(payload).encode("utf-8"),
        method="POST",
        headers={
            "Content-Type": "application/json",
            "Authorization": "Basic " + base64.b64encode(auth_raw).decode("ascii"),
        },
    )
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            data = resp.read().decode("utf-8", errors="replace")
            if 200 <= resp.status < 300:
                return "sent", ""
            return "failed", f"Mailjet API returned HTTP {resp.status}: {data[:500]}"
    except urllib.error.HTTPError as exc:
        body = exc.read().decode("utf-8", errors="replace")[:700]
        if exc.code in {401, 403}:
            return "failed", "Mailjet API authentication failed. Check MAILJET_API_KEY and MAILJET_SECRET_KEY."
        return "failed", f"Mailjet API returned HTTP {exc.code}: {body}"
    except Exception as exc:
        return "failed", f"Mailjet API send failed: {str(exc)[:700]}"


def send_outreach_email(recipient_email: str, subject: str, body: str, recipient_name: str = "") -> Tuple[str, str]:
    """Send through Mailjet API when configured, otherwise through SMTP. If not configured, log as queued."""
    provider = EMAIL_PROVIDER or "smtp"
    if provider in {"mailjet", "mailjet_api"}:
        return send_mailjet_api_email(recipient_email, recipient_name, subject, body)

    if not smtp_is_configured():
        return "queued_smtp_not_configured", "SMTP not configured. Set SMTP_HOST, SMTP_PORT, SMTP_FROM_EMAIL and optional SMTP_USERNAME/SMTP_PASSWORD."

    msg = EmailMessage()
    msg["From"] = f"{SMTP_FROM_NAME} <{SMTP_FROM_EMAIL}>"
    msg["To"] = recipient_email
    msg["Subject"] = subject
    msg.set_content(body)

    try:
        socket.getaddrinfo(SMTP_HOST, SMTP_PORT)
        with smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=20) as smtp:
            if SMTP_USE_TLS:
                smtp.starttls()
            if SMTP_USERNAME and SMTP_PASSWORD:
                smtp.login(SMTP_USERNAME, SMTP_PASSWORD)
            smtp.send_message(msg)
        return "sent", ""
    except Exception as exc:
        return "failed", friendly_smtp_error(exc)


def get_current_user(authorization: str = Header(default="")) -> Dict[str, Any]:
    if not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing login token")
    token = authorization.replace("Bearer ", "", 1).strip()
    if len(token) < 20:
        raise HTTPException(status_code=401, detail="Invalid session")
    with get_db() as conn:
        row = conn.execute(
            """
            SELECT users.id, users.username, users.full_name, users.email, users.phone, users.photo_file_name, users.photo_file_path,
                   users.role, users.is_active, users.force_password_change, sessions.created_at, sessions.expires_at
            FROM sessions JOIN users ON sessions.user_id = users.id
            WHERE sessions.token = ?
            """,
            (token,),
        ).fetchone()
        if not row or row["is_active"] != 1:
            raise HTTPException(status_code=401, detail="Invalid or inactive session")
        expires = parse_optional_iso(row["expires_at"]) or (parse_iso(row["created_at"]) + timedelta(hours=SESSION_TTL_HOURS))
        if expires < datetime.utcnow():
            conn.execute("DELETE FROM sessions WHERE token=?", (token,))
            conn.commit()
            raise HTTPException(status_code=401, detail="Session expired. Please login again.")
        conn.execute("UPDATE sessions SET last_seen_at=? WHERE token=?", (now_iso(), token))
        conn.commit()
    data = dict_row(row)
    data.pop("created_at", None)
    data.pop("expires_at", None)
    return data


def require_roles(*roles: str):
    def _checker(user: Dict[str, Any] = Depends(get_current_user)) -> Dict[str, Any]:
        if user["role"] not in roles:
            raise HTTPException(status_code=403, detail="Insufficient permission")
        return user
    return _checker


def compute_total(a: AssessmentCreate) -> int:
    return (
        a.technical_score + a.project_score + a.practical_score + a.communication_score
        + a.client_readiness_score + a.cost_fitment_score + a.availability_score
    )


def status_from_score(score: int) -> str:
    if score >= 85:
        return "A1 - Ready to Deploy"
    if score >= 75:
        return "A2 - Deployable in 15 Days"
    if score >= 65:
        return "B - Keep Warm"
    return "Rejected / Archive"


def rating_level(score: int) -> str:
    if score >= 85:
        return "Excellent Fit"
    if score >= 75:
        return "Strong Fit"
    if score >= 65:
        return "Good Fit"
    if score >= 50:
        return "Needs Review"
    return "Low Fit"


def risk_level(score: int) -> str:
    if score >= 75:
        return "High Risk"
    if score >= 45:
        return "Medium Risk"
    if score >= 20:
        return "Low Risk"
    return "Clean / Low Signal"


def fake_risk_rag(score: Any) -> Dict[str, str]:
    value = int(score or 0)
    if value >= 75:
        return {"rag": "Red", "label": "Red - High Fake Risk", "action": "Do not submit without manual validation"}
    if value >= 45:
        return {"rag": "Amber", "label": "Amber - Needs Review", "action": "Review risk reasons before shortlisting"}
    return {"rag": "Green", "label": "Green - Low Signal", "action": "Proceed with normal screening"}


def enrich_fake_risk_rag(item: Dict[str, Any]) -> Dict[str, Any]:
    rag = fake_risk_rag(item.get("fake_risk_score", 0))
    item["fake_risk_rag"] = rag["rag"]
    item["fake_risk_rag_label"] = rag["label"]
    item["fake_risk_rag_action"] = rag["action"]
    return item


def generate_candidate_code(conn: sqlite3.Connection) -> str:
    year = datetime.utcnow().year
    row = conn.execute("SELECT COUNT(*) AS c FROM candidates").fetchone()
    return f"TRX-RB-{year}-{int(row['c']) + 1:04d}"


def generate_demand_code(conn: sqlite3.Connection) -> str:
    year = datetime.utcnow().year
    row = conn.execute("SELECT COUNT(*) AS c FROM demand_requests").fetchone()
    return f"TRX-DMD-{year}-{int(row['c']) + 1:04d}"


def safe_date(value: Any) -> Optional[datetime]:
    text = str(value or "").strip()
    if not text or text.lower() in ["immediate", "to be confirmed", "tbd", "na", "n/a"]:
        return None
    try:
        return datetime.fromisoformat(text.replace("Z", "").split("T")[0])
    except Exception:
        return None


def availability_fit(candidate: Dict[str, Any], demand: Dict[str, Any]) -> Tuple[int, str]:
    status = str(candidate.get("current_status") or "").lower()
    notice = int(candidate.get("notice_period_days") or 0)
    available_by = safe_date(candidate.get("available_by_date"))
    demand_start = safe_date(demand.get("start_date"))
    if status in ["available", "bench", "freelance"] and notice <= 7:
        return 10, "Immediate / available"
    if demand_start and available_by:
        if available_by <= demand_start:
            return 9, "Available by demand start date"
        delta = (available_by - demand_start).days
        if delta <= 14:
            return 5, f"Available {delta} day(s) after demand start"
        return -2, f"Available {delta} day(s) after demand start"
    if status == "notice period":
        if notice <= 15:
            return 6, "Short notice period"
        if notice <= 30:
            return 4, "Standard notice period"
        if notice <= 60:
            return 1, "Long notice period"
        return -3, "Very long notice period"
    if status in ["not available", "employed"] and notice > 45:
        return -2, "Availability risk"
    return 2, "Availability to be confirmed"


def level_from_match(score: int) -> str:
    if score >= 85:
        return "Excellent Match"
    if score >= 75:
        return "Strong Match"
    if score >= 60:
        return "Review Match"
    if score >= 40:
        return "Weak Match"
    return "Not Suitable"


def demand_required_skills(demand: Dict[str, Any]) -> List[str]:
    explicit = extract_skills(str(demand.get("required_skills") or ""))
    contextual = extract_skills(" ".join([
        str(demand.get("role_title") or ""),
        str(demand.get("role_definition") or ""),
    ]))
    combined = []
    for skill in explicit + contextual:
        if skill not in combined:
            combined.append(skill)
    return combined


def candidate_available_skills(candidate: Dict[str, Any]) -> List[str]:
    text = " ".join([
        str(candidate.get("primary_skill") or ""),
        str(candidate.get("secondary_skills") or ""),
        str(candidate.get("skill_matches") or ""),
        str(candidate.get("resume_text") or ""),
        str(candidate.get("domain_exposure") or ""),
        str(candidate.get("current_company") or ""),
        str(candidate.get("previous_companies") or ""),
        str(candidate.get("project_details") or ""),
    ])
    return extract_skills(text)


def role_title_alignment(candidate: Dict[str, Any], demand: Dict[str, Any], matches: List[str]) -> Tuple[int, str]:
    title_text = str(demand.get("role_title") or "").lower()
    candidate_text = " ".join([
        str(candidate.get("primary_skill") or ""),
        str(candidate.get("secondary_skills") or ""),
        str(candidate.get("resume_text") or ""),
        str(candidate.get("project_details") or ""),
    ]).lower()
    title_skills = extract_skills(title_text)
    if title_skills and any(s in matches for s in title_skills):
        return 10, "Role-title skill matched"
    if any(word in title_text for word in ["lead", "architect", "manager", "senior"]):
        exp = float(candidate.get("relevant_experience") or candidate.get("total_experience") or 0)
        if exp < 5:
            return -6, "Seniority risk"
    if title_text and any(word in candidate_text for word in [w for w in re.split(r"[^a-z0-9]+", title_text) if len(w) > 3]):
        return 4, "Role-title wording partially matched"
    return 0, "No explicit title alignment"


def domain_alignment(candidate: Dict[str, Any], demand: Dict[str, Any]) -> Tuple[int, str]:
    demand_domains = extract_domains(" ".join([str(demand.get("domain") or ""), str(demand.get("role_definition") or "")]))
    candidate_domains = extract_domains(" ".join([str(candidate.get("domain_exposure") or ""), str(candidate.get("resume_text") or ""), str(candidate.get("project_details") or "")]))
    if not demand_domains:
        return 2, "Demand domain not specified"
    common = [d for d in demand_domains if d in candidate_domains]
    if common:
        return 8, "Domain matched: " + ", ".join(common)
    return -4, "Domain not evident"


def experience_alignment(candidate: Dict[str, Any], demand: Dict[str, Any]) -> Tuple[int, str]:
    exp = float(candidate.get("relevant_experience") or candidate.get("total_experience") or 0)
    text = " ".join([str(demand.get("role_title") or ""), str(demand.get("role_definition") or "")]).lower()
    min_exp = 0.0
    exp_matches = re.findall(r"(\d{1,2}(?:\.\d+)?)\s*(?:\+)?\s*(?:years|yrs|year)", text)
    if exp_matches:
        min_exp = max(float(x) for x in exp_matches)
    elif any(w in text for w in ["architect", "lead", "senior", "manager"]):
        min_exp = 5.0
    elif any(w in text for w in ["junior", "associate", "fresher"]):
        min_exp = 0.0
    if min_exp and exp < min_exp:
        return -8, f"Experience below role expectation: {exp:g}/{min_exp:g} yrs"
    if exp >= max(min_exp, 3):
        return 10, f"Experience fit: {exp:g} yrs"
    if exp > 0:
        return 5, f"Limited experience: {exp:g} yrs"
    return -6, "Experience not evident"


def strict_match_notes(required: List[str], matches: List[str], gaps: List[str], score: int) -> str:
    if required and not matches:
        return "Rejected by strict screen: no required skills found."
    if required and len(matches) / max(1, len(required)) < 0.4:
        return "Strict screen: low required-skill coverage."
    if score >= 75:
        return "Strong skill and fit evidence."
    if score >= 60:
        return "Usable shortlist candidate; review gaps."
    if score >= 40:
        return "Weak fit; shortlist only if bench options are limited."
    return "Not suitable for this role."


def match_candidate_to_demand(candidate: Dict[str, Any], demand: Dict[str, Any]) -> Dict[str, Any]:
    required = demand_required_skills(demand)
    available = candidate_available_skills(candidate)
    matches = [s for s in required if s in available]
    gaps = [s for s in required if s not in available]
    match_ratio = len(matches) / max(1, len(required)) if required else 0.0

    matched_clusters = {skill_cluster(s) for s in matches}
    required_clusters = {skill_cluster(s) for s in required}
    cluster_ratio = len(matched_clusters) / max(1, len(required_clusters)) if required_clusters else 0.0

    title_bonus, title_note = role_title_alignment(candidate, demand, matches)
    domain_bonus, domain_note = domain_alignment(candidate, demand)
    exp_bonus, exp_note = experience_alignment(candidate, demand)
    availability_bonus, availability_fit_text = availability_fit(candidate, demand)

    ml = int(candidate.get("ml_rating_score") or 0)
    fake_risk = int(candidate.get("fake_risk_score") or 0)
    rate = float(candidate.get("negotiated_rate") or candidate.get("expected_rate") or 0)
    max_cost = float(demand.get("max_internal_cost") or 0)
    commercial_bonus = 5 if max_cost and rate and rate <= max_cost else (2 if not max_cost or not rate else -6)
    commercial_fit = "Within cost band" if max_cost and rate and rate <= max_cost else ("Cost not available" if not rate else "Above cost band")

    # Stricter weighting: required skills dominate. Experience/availability/commercials cannot rescue a poor skill fit.
    skill_score = match_ratio * 58
    cluster_score = cluster_ratio * 8
    resume_quality_score = min(max(ml, 0), 100) * 0.05 if ml else 0
    risk_penalty = min(14, fake_risk * 0.14)

    raw = skill_score + cluster_score + title_bonus + domain_bonus + exp_bonus + min(8, availability_bonus) + commercial_bonus + resume_quality_score - risk_penalty

    if required and not matches:
        raw = min(raw, 34)
    elif required and match_ratio < 0.40:
        raw = min(raw, 58)
    elif required and match_ratio < 0.60:
        raw = min(raw, 70)
    if fake_risk >= 75:
        raw = min(raw, 55)

    score = max(0, min(100, int(round(raw))))
    notes = strict_match_notes(required, matches, gaps, score)
    return {
        "match_score": score,
        "match_level": level_from_match(score),
        "required_skills": required,
        "available_skills": available,
        "skill_matches": matches,
        "skill_gaps": gaps,
        "skill_match_ratio": round(match_ratio, 2),
        "commercial_fit": commercial_fit,
        "availability_fit": availability_fit_text,
        "title_fit": title_note,
        "domain_fit": domain_note,
        "experience_fit": exp_note,
        "strict_notes": notes,
    }


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or " ").strip()


def clean_extracted_text(text: str) -> str:
    value = text or ""
    value = value.replace("\x00", " ")
    value = re.sub(r"[\u200b\u200c\u200d\ufeff]", "", value)
    value = re.sub(r"(?<=[a-z])(?=[A-Z])", " ", value)
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def canonicalize_name(name: str) -> str:
    parts = []
    for word in re.split(r"\s+", re.sub(r"[^A-Za-z .'-]", " ", name or "").strip()):
        if not word:
            continue
        if word.isupper():
            parts.append(word.title())
        else:
            parts.append(word[0].upper() + word[1:])
    return " ".join(parts).strip()


def extract_skills(text: str) -> List[str]:
    lower = f" {normalize_text(text).lower()} "
    found = []
    for canonical, aliases in SKILL_ALIASES.items():
        for alias in aliases:
            pattern = r"(?<![a-z0-9])" + re.escape(alias.lower()) + r"(?![a-z0-9])"
            if re.search(pattern, lower):
                found.append(canonical)
                break
    return sorted(set(found))


def extract_domains(text: str) -> List[str]:
    lower = normalize_text(text).lower()
    found = []
    for d in DOMAIN_TERMS:
        if d.lower() in lower and d not in found:
            found.append(d)
    return found


def years_between(start_year: int, end_year: int) -> float:
    if start_year < 1970 or end_year < 1970 or end_year < start_year:
        return 0.0
    return min(50.0, float(end_year - start_year))


def extract_experience_years(text: str) -> float:
    lower = normalize_text(text).lower()
    values: List[float] = []

    # Direct mentions: 8.8 years, 7 yrs, 10+ year experience.
    for match in re.finditer(r"(\d{1,2}(?:\.\d+)?)\s*(?:\+)?\s*(?:years?|yrs?|yr)\b", lower):
        try:
            val = float(match.group(1))
            if 0 <= val <= 45:
                values.append(val)
        except ValueError:
            pass

    # Text like "Total Experience: 6.6 years". Keep this strict so "Experience 2018 - Present"
    # does not get incorrectly decoded as 20 years.
    for match in re.finditer(r"(?:total|overall|professional)\s+experience\s*[:\-]?\s*(\d{1,2}(?:\.\d+)?)(?!\d)", lower):
        try:
            val = float(match.group(1))
            if 0 <= val <= 45:
                values.append(val)
        except ValueError:
            pass
    for match in re.finditer(r"\bexperience\s*[:\-]\s*(\d{1,2}(?:\.\d+)?)(?!\d)", lower):
        try:
            val = float(match.group(1))
            if 0 <= val <= 45:
                values.append(val)
        except ValueError:
            pass

    # Date ranges in work history: 2018 - 2024, Mar 2019 to Present.
    current_year = datetime.utcnow().year
    for match in re.finditer(r"\b(19\d{2}|20\d{2})\b\s*(?:-|–|—|to|till|until)\s*(present|current|now|19\d{2}|20\d{2})\b", lower):
        start_year = int(match.group(1))
        end_raw = match.group(2)
        end_year = current_year if end_raw in {"present", "current", "now"} else int(end_raw)
        span = years_between(start_year, end_year)
        if span:
            values.append(span)

    # Month-year ranges: Jan 2020 - Mar 2023.
    month_year_ranges = re.findall(
        r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\s+(19\d{2}|20\d{2})\s*(?:-|–|—|to)\s*(?:(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\s+)?(present|current|now|19\d{2}|20\d{2})",
        lower,
    )
    for start_y, end_y_raw in month_year_ranges:
        end_y = current_year if end_y_raw in {"present", "current", "now"} else int(end_y_raw)
        span = years_between(int(start_y), end_y)
        if span:
            values.append(span)

    # If no direct experience but multiple job years are present, infer timeline span conservatively.
    years = sorted({int(y) for y in re.findall(r"\b(19\d{2}|20\d{2})\b", lower)})
    if len(years) >= 2:
        inferred = years_between(years[0], min(years[-1], current_year))
        if inferred:
            values.append(inferred)

    if not values:
        return 0.0
    # Prefer explicit direct values; use the highest credible value, rounded to one decimal.
    return round(max(values), 1)


def extract_section(text: str, headings: List[str], max_chars: int = 1500) -> str:
    pattern = r"(?is)(?:^|\n)\s*(?:" + "|".join(re.escape(h) for h in headings) + r")\s*[:\-]?\s*\n?(.*?)(?=\n\s*(?:summary|profile|objective|skills|technical skills|experience|employment|work history|education|certifications|projects|personal details|contact)\s*[:\-]?\s*\n|$)"
    m = re.search(pattern, text)
    return (m.group(1).strip()[:max_chars] if m else "")


def extract_contact_and_name(text: str) -> Dict[str, str]:
    cleaned_text = clean_extracted_text(text)
    email_match = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", cleaned_text)
    phone_match = re.search(r"(?:\+?\d[\d\s().-]{7,}\d)", cleaned_text)
    lines = [l.strip(" •|\t-") for l in cleaned_text.splitlines() if l.strip(" •|\t-")]

    bad_fragments = {
        "resume", "curriculum vitae", "cv", "profile", "email", "phone", "mobile", "objective", "summary",
        "experience", "skills", "education", "linkedin", "github", "address", "contact", "professional"
    }
    name = ""

    # Prefer explicit labels: Name: John Doe.
    for line in lines[:20]:
        m = re.search(r"\b(?:name|candidate name|full name)\s*[:\-]\s*([A-Za-z][A-Za-z .'-]{2,80})", line, flags=re.I)
        if m:
            candidate = canonicalize_name(m.group(1))
            if 2 <= len(candidate.split()) <= 5:
                name = candidate
                break

    # Otherwise use the first clean human-looking line near the top, avoiding contact and headings.
    if not name:
        for line in lines[:18]:
            if "@" in line or re.search(r"https?://|www\.|linkedin|github|\+?\d[\d\s().-]{7,}\d", line, flags=re.I):
                continue
            cleaned = re.sub(r"[^A-Za-z .'-]", " ", line).strip()
            cleaned = re.sub(r"\s+", " ", cleaned)
            words = cleaned.split()
            lower = cleaned.lower()
            if not cleaned or any(b in lower for b in bad_fragments):
                continue
            if 2 <= len(words) <= 5 and all(len(w) > 1 or w.upper() in {"S", "R", "B", "K"} for w in words):
                # Avoid lines that are clearly job titles.
                if not re.search(r"\b(developer|engineer|manager|consultant|analyst|architect|administrator|specialist|lead|intern)\b", lower):
                    name = canonicalize_name(cleaned)
                    break

    phone = phone_match.group(0).strip() if phone_match else ""
    phone = re.sub(r"\s+", " ", phone)
    return {
        "full_name": name,
        "email": email_match.group(0).lower() if email_match else "",
        "phone": phone,
    }


def extract_text_from_file(path: str, filename: str) -> str:
    ext = os.path.splitext(filename.lower())[1]
    if ext in [".txt", ".md", ".csv"]:
        with open(path, "rb") as f:
            raw = f.read()
        return clean_extracted_text(raw.decode("utf-8", errors="ignore"))
    if ext == ".pdf":
        try:
            from pypdf import PdfReader  # type: ignore
            reader = PdfReader(path)
            parts = []
            for page in reader.pages:
                parts.append(page.extract_text() or "")
            return clean_extracted_text("\n".join(parts))
        except Exception as exc:
            return f"PDF text extraction unavailable. File name: {filename}. Error: {exc}"
    if ext == ".docx":
        try:
            from docx import Document  # type: ignore
            doc = Document(path)
            parts = []
            for p in doc.paragraphs:
                if p.text:
                    parts.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    vals = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if vals:
                        parts.append(" | ".join(vals))
            return clean_extracted_text("\n".join(parts))
        except Exception as exc:
            return f"DOCX text extraction unavailable. File name: {filename}. Error: {exc}"
    with open(path, "rb") as f:
        raw = f.read(250000)
    return clean_extracted_text(raw.decode("utf-8", errors="ignore"))


def build_standard_summary(name: str, skills: List[str], domains: List[str], exp: float, role_title: str, fit_score: int, fake_score: int) -> str:
    primary = skills[0] if skills else "Unspecified"
    skill_text = ", ".join(skills) if skills else "Not clearly detected"
    domain_text = ", ".join(domains) if domains else "Not clearly detected"
    return (
        f"Candidate: {name or 'Name not detected'}\n"
        f"Target Role: {role_title or 'Not specified'}\n"
        f"Primary Skill: {primary}\n"
        f"Detected Skills: {skill_text}\n"
        f"Domain Exposure: {domain_text}\n"
        f"Detected Experience: {exp:g} years\n"
        f"ML Resume Fit Score: {fit_score}/100\n"
        f"Fake Resume Risk Score: {fake_score}/100\n"
        f"Generated by 1Resource Resume Intelligence based on uploaded resume and role definition."
    )


def analyze_resume_text(resume_text: str, role_definition: str = "", role_title: str = "") -> Dict[str, Any]:
    text = resume_text or ""
    normalized = normalize_text(text)
    skills = extract_skills(text)
    role_skills = extract_skills(role_definition + " " + role_title)
    domains = extract_domains(text)
    exp = extract_experience_years(text)
    contact = extract_contact_and_name(text)

    if role_skills:
        matches = [s for s in role_skills if s in skills]
        gaps = [s for s in role_skills if s not in skills]
        match_ratio = len(matches) / max(1, len(role_skills))
    else:
        matches = skills[:]
        gaps = []
        match_ratio = min(1.0, len(skills) / 6)

    score = int(min(100, round(match_ratio * 58 + min(exp, 12) * 2.2 + min(len(skills), 10) * 2 + (8 if domains else 0))))
    if contact.get("email") and contact.get("phone"):
        score = min(100, score + 5)
    if role_definition and len(role_definition) > 40 and not role_skills:
        score = max(0, score - 8)

    risk = 0
    reasons: List[str] = []
    if len(normalized) < 700:
        risk += 20
        reasons.append("Resume text is too short for a reliable evaluation.")
    if not contact.get("email"):
        risk += 12
        reasons.append("Email address not detected.")
    if not contact.get("phone"):
        risk += 10
        reasons.append("Phone number not detected.")
    if exp >= 8 and len(re.findall(r"\b(20\d{2}|19\d{2})\b", text)) < 2:
        risk += 18
        reasons.append("High experience claimed but limited timeline/year evidence found.")
    if exp >= 6 and len(re.findall(r"\b(project|client|company|employer|implementation|deployed|delivered)\b", text.lower())) < 4:
        risk += 14
        reasons.append("Experience claim is not sufficiently supported by project or employer evidence.")
    buzz_terms = re.findall(r"\b(expert|expertise|end[- ]to[- ]end|responsible for|handled|worked on|proficient|dynamic|hardworking)\b", text.lower())
    evidence_terms = re.findall(r"\b(api|database|deployment|architecture|testing|integration|module|performance|security|cloud|production)\b", text.lower())
    if len(buzz_terms) > 16 and len(evidence_terms) < 8:
        risk += 16
        reasons.append("High number of generic claims with limited technical evidence.")
    if role_skills and not matches:
        risk += 20
        reasons.append("No clear match found against the role definition skills.")
    if re.search(r"lorem ipsum|sample resume|dummy candidate|fake resume", text.lower()):
        risk += 30
        reasons.append("Placeholder or dummy-resume wording detected.")
    if len(set(re.findall(r"\b[A-Za-z]{5,}\b", text.lower()))) < 80 and len(normalized) > 1200:
        risk += 10
        reasons.append("Text appears repetitive with low vocabulary diversity.")
    risk = min(100, risk)
    if not reasons:
        reasons.append("No major fake-resume signal detected by the local screening model.")

    return {
        "detected_name": contact.get("full_name", ""),
        "email": contact.get("email", ""),
        "phone": contact.get("phone", ""),
        "skills": skills,
        "role_skills": role_skills,
        "skill_matches": matches,
        "skill_gaps": gaps,
        "domains": domains,
        "experience_years": exp,
        "fit_score": score,
        "rating_level": rating_level(score),
        "fake_risk_score": risk,
        "fake_risk_level": risk_level(risk),
        "fake_risk_reasons": reasons,
    }


def sanitize_filename(name: str) -> str:
    base = os.path.basename(name or "resume.txt")
    base = re.sub(r"[^A-Za-z0-9._ -]", "_", base).strip(" .")
    return base[:120] or "resume.txt"


def validate_upload_file(file: UploadFile) -> str:
    safe_name = sanitize_filename(file.filename or "resume.txt")
    ext = os.path.splitext(safe_name.lower())[1]
    if ext not in ALLOWED_UPLOAD_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"Unsupported file type. Allowed: {', '.join(sorted(ALLOWED_UPLOAD_EXTENSIONS))}")
    return safe_name


def validate_photo_file(file: UploadFile) -> str:
    safe_name = sanitize_filename(file.filename or "photograph.jpg")
    ext = os.path.splitext(safe_name.lower())[1]
    if ext not in ALLOWED_PHOTO_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"Unsupported photograph type. Allowed: {', '.join(sorted(ALLOWED_PHOTO_EXTENSIONS))}")
    return safe_name


def save_photo_upload(file: UploadFile, prefix: str) -> Tuple[str, str]:
    safe_name = validate_photo_file(file)
    stored_name = f"{prefix}_{uuid.uuid4().hex}_{safe_name}"
    path = os.path.abspath(os.path.join(UPLOAD_DIR, stored_name))
    if not path.startswith(os.path.abspath(UPLOAD_DIR) + os.sep):
        raise HTTPException(status_code=400, detail="Invalid photograph upload path")
    size = 0
    with open(path, "wb") as f:
        while True:
            chunk = file.file.read(1024 * 1024)
            if not chunk:
                break
            size += len(chunk)
            if size > MAX_PHOTO_BYTES:
                f.close()
                try:
                    os.remove(path)
                except OSError:
                    pass
                raise HTTPException(status_code=413, detail=f"Photograph is too large. Maximum allowed size is {MAX_PHOTO_MB} MB")
            f.write(chunk)
    if size == 0:
        try:
            os.remove(path)
        except OSError:
            pass
        raise HTTPException(status_code=400, detail="Uploaded photograph is empty")
    return safe_name, path


def save_upload(file: UploadFile, prefix: str) -> Tuple[str, str]:
    safe_name = validate_upload_file(file)
    stored_name = f"{prefix}_{uuid.uuid4().hex}_{safe_name}"
    path = os.path.abspath(os.path.join(UPLOAD_DIR, stored_name))
    if not path.startswith(os.path.abspath(UPLOAD_DIR) + os.sep):
        raise HTTPException(status_code=400, detail="Invalid upload path")
    size = 0
    with open(path, "wb") as f:
        while True:
            chunk = file.file.read(1024 * 1024)
            if not chunk:
                break
            size += len(chunk)
            if size > MAX_UPLOAD_BYTES:
                f.close()
                try:
                    os.remove(path)
                except OSError:
                    pass
                raise HTTPException(status_code=413, detail=f"Resume file is too large. Maximum allowed size is {MAX_UPLOAD_MB} MB")
            f.write(chunk)
    if size == 0:
        try:
            os.remove(path)
        except OSError:
            pass
        raise HTTPException(status_code=400, detail="Uploaded file is empty")
    return safe_name, path


def ensure_safe_file_path(path: str) -> str:
    abs_path = os.path.abspath(path or "")
    if not abs_path.startswith(os.path.abspath(UPLOAD_DIR) + os.sep):
        raise HTTPException(status_code=403, detail="File path is outside the upload repository")
    if not os.path.exists(abs_path):
        raise HTTPException(status_code=404, detail="Resume file not found")
    return abs_path


def create_resume_record(
    conn: sqlite3.Connection,
    candidate_id: int,
    file_name: str,
    file_path: str,
    resume_text: str,
    role_title: str,
    role_definition: str,
    source: str,
    analysis: Dict[str, Any],
    update_identity: bool = True,
) -> int:
    cur = conn.execute(
        """
        INSERT INTO resumes(candidate_id, role_title, role_definition, resume_file_name, resume_file_path, resume_text,
            extracted_json, fit_score, rating_level, fake_risk_score, fake_risk_level, fake_risk_reasons,
            skill_matches, skill_gaps, source, created_at)
        VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
        """,
        (
            candidate_id, role_title, role_definition, file_name, file_path, resume_text,
            json.dumps(analysis), int(analysis["fit_score"]), analysis["rating_level"],
            int(analysis["fake_risk_score"]), analysis["fake_risk_level"],
            " | ".join(analysis["fake_risk_reasons"]), ", ".join(analysis["skill_matches"]),
            ", ".join(analysis["skill_gaps"]), source, now_iso()
        ),
    )
    resume_id = int(cur.lastrowid)
    all_skills = analysis.get("skills", [])
    primary_skill = all_skills[0] if all_skills else ""
    secondary = ", ".join(all_skills[1:]) if len(all_skills) > 1 else ""
    domain = ", ".join(analysis.get("domains", []))
    summary = build_standard_summary(
        analysis.get("detected_name", ""), all_skills, analysis.get("domains", []),
        float(analysis.get("experience_years") or 0), role_title, int(analysis["fit_score"]), int(analysis["fake_risk_score"])
    )
    existing = conn.execute("SELECT * FROM candidates WHERE id=?", (candidate_id,)).fetchone()
    full_name = existing["full_name"] if existing and existing["full_name"] else analysis.get("detected_name") or "Unknown Candidate"
    if update_identity and full_name in ["Unknown Candidate", "Candidate"] and analysis.get("detected_name"):
        full_name = analysis["detected_name"]
    if not update_identity and existing:
        full_name = existing["full_name"]
    email = (existing["email"] if existing else "") or (analysis.get("email", "") if update_identity else "")
    phone = (existing["phone"] if existing else "") or (analysis.get("phone", "") if update_identity else "")
    exp = existing["total_experience"] or analysis.get("experience_years", 0)
    relevant_exp = existing["relevant_experience"] or min(float(exp or 0), float(analysis.get("experience_years") or 0))
    conn.execute(
        """
        UPDATE candidates SET
            full_name=?, email=?, phone=?, total_experience=?, relevant_experience=?,
            primary_skill=COALESCE(NULLIF(?,''), primary_skill),
            secondary_skills=COALESCE(NULLIF(?,''), secondary_skills),
            domain_exposure=COALESCE(NULLIF(?,''), domain_exposure),
            resume_text=?, resume_file_name=?, resume_file_path=?, last_role_definition=?,
            ml_rating_score=?, ml_rating_level=?, fake_risk_score=?, fake_risk_level=?, fake_risk_reasons=?,
            skill_matches=?, skill_gaps=?, resume_count=(SELECT COUNT(*) FROM resumes WHERE candidate_id=?), updated_at=?
        WHERE id=?
        """,
        (
            full_name, email, phone, exp, relevant_exp, primary_skill, secondary, domain,
            summary, file_name, file_path, role_definition, int(analysis["fit_score"]), analysis["rating_level"],
            int(analysis["fake_risk_score"]), analysis["fake_risk_level"], " | ".join(analysis["fake_risk_reasons"]),
            ", ".join(analysis["skill_matches"]), ", ".join(analysis["skill_gaps"]), candidate_id, now_iso(), candidate_id
        ),
    )
    return resume_id


def clean_resume_value(value: Any, fallback: str = "-") -> str:
    text = str(value or "").strip()
    return text if text else fallback


def build_standardized_resume(candidate: Dict[str, Any]) -> str:
    full_name = clean_resume_value(candidate.get("full_name"), "Candidate")
    candidate_code = clean_resume_value(candidate.get("candidate_code"))
    role = clean_resume_value(candidate.get("primary_skill"), "Professional")
    secondary = clean_resume_value(candidate.get("secondary_skills"))
    company = clean_resume_value(candidate.get("current_company"), "Not specified")
    previous = clean_resume_value(candidate.get("previous_companies"), "Not specified")
    projects = clean_resume_value(candidate.get("project_details") or candidate.get("resume_text"), "Project details not available")
    skills = ", ".join([x for x in [candidate.get("primary_skill"), candidate.get("secondary_skills")] if x]) or "Not specified"
    lines = [
        "1RESOURCE STANDARDIZED RESUME",
        "Generated by Truflux Technologies",
        "=" * 72,
        "",
        "CANDIDATE DETAILS",
        f"Name: {full_name}",
        f"Candidate Code: {candidate_code}",
        f"Email: {clean_resume_value(candidate.get('email'))}",
        f"Phone: {clean_resume_value(candidate.get('phone'))}",
        f"Location: {clean_resume_value(candidate.get('location'))}",
        f"Availability: {clean_resume_value(candidate.get('current_status'))}",
        f"Available By: {clean_resume_value(candidate.get('available_by_date'))}",
        f"Notice Period: {clean_resume_value(candidate.get('notice_period_days'), '0')} days",
        "",
        "PROFESSIONAL POSITIONING",
        f"Target / Primary Role: {role}",
        f"Internal Level: {clean_resume_value(candidate.get('internal_level'))}",
        f"Total Experience: {clean_resume_value(candidate.get('total_experience'), '0')} years",
        f"Relevant Experience: {clean_resume_value(candidate.get('relevant_experience'), '0')} years",
        f"Proficiency: {clean_resume_value(candidate.get('proficiency'))}",
        "",
        "SKILLS",
        f"Primary Skill: {clean_resume_value(candidate.get('primary_skill'))}",
        f"Secondary Skills: {secondary}",
        f"Detected / Matched Skills: {clean_resume_value(candidate.get('skill_matches'))}",
        f"Skill Gaps: {clean_resume_value(candidate.get('skill_gaps'))}",
        "",
        "COMPANY / PROJECT DETAILS",
        f"Current / Last Company: {company}",
        f"Previous Companies: {previous}",
        "Project / Assignment Details:",
        projects,
        "",
        "DOMAIN AND CERTIFICATIONS",
        f"Domain Exposure: {clean_resume_value(candidate.get('domain_exposure'))}",
        f"Certifications: {clean_resume_value(candidate.get('certifications'))}",
        f"Portfolio URL: {clean_resume_value(candidate.get('portfolio_url'))}",
        "",
        "1RESOURCE SCREENING SUMMARY",
        f"Resume Fit Score: {clean_resume_value(candidate.get('ml_rating_score'), '0')}/100",
        f"Resume Fit Level: {clean_resume_value(candidate.get('ml_rating_level'), 'Not rated')}",
        f"Fake Resume Risk Score: {clean_resume_value(candidate.get('fake_risk_score'), '0')}/100",
        f"Fake Resume Risk Level: {clean_resume_value(candidate.get('fake_risk_level'), 'Not checked')}",
        f"Fake Risk RAG: {clean_resume_value(candidate.get('fake_risk_rag'), 'Not available')}",
        f"Risk Reasons: {clean_resume_value(candidate.get('fake_risk_reasons'))}",
        "",
        "STANDARDIZED SUMMARY",
        clean_resume_value(candidate.get("resume_text"), "Resume summary not available."),
        "",
        "COMMERCIAL DETAILS",
        f"Expected Rate: {clean_resume_value(candidate.get('expected_rate'), '0')}",
        f"Negotiated Rate: {clean_resume_value(candidate.get('negotiated_rate'), '0')}",
        "",
        "DOCUMENT CONTROL",
        f"Generated On: {now_iso()}",
        "Format: 1Resource standardized resume format",
        "Prepared by: Truflux Technologies",
    ]
    return "\n".join(lines)


def pdf_text(value: Any, fallback: str = "-") -> str:
    return escape(clean_resume_value(value, fallback)).replace("\n", "<br/>")


CONTACT_SUPPRESSED = "Suppressed - candidate contact protected"
CONTACT_NOT_CONFIGURED = "Not configured in login profile"


def authorized_contact_value(value: Any) -> str:
    value = str(value or "").strip()
    return value if value else CONTACT_NOT_CONFIGURED


def redact_contact_info(value: Any) -> str:
    text = str(value or "")
    if not text:
        return ""
    # Redact email addresses.
    text = re.sub(r"\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b", "[email suppressed]", text)
    # Redact phone numbers while avoiding ordinary short numbers and scores.
    text = re.sub(r"(?<!\w)(?:\+?\d[\d\s().-]{7,}\d)(?!\w)", "[phone suppressed]", text)
    return text


def pdf_text_redacted(value: Any, fallback: str = "-") -> str:
    return escape(clean_resume_value(redact_contact_info(value), fallback)).replace("\n", "<br/>")


def compact_resume_text(value: Any, limit: int = 2200) -> str:
    text = re.sub(r"\s+", " ", redact_contact_info(value)).strip()
    if not text:
        return "-"
    return text if len(text) <= limit else text[:limit].rstrip() + " ... [truncated for PDF readability]"


def parse_resume_analysis_json(value: Any) -> Dict[str, Any]:
    try:
        if isinstance(value, dict):
            return value
        if value:
            return json.loads(str(value))
    except Exception:
        return {}
    return {}


def make_kv_table(rows: List[Tuple[str, Any]], styles: Dict[str, Any]) -> Table:
    data = [[Paragraph(f"<b>{escape(label)}</b>", styles["small"]), Paragraph(pdf_text(value), styles["body"])] for label, value in rows]
    table = Table(data, colWidths=[1.65 * inch, 4.55 * inch])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f1f5f9")),
        ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#0f172a")),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#dbe4f0")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    return table


def section_title(text_value: str, styles: Dict[str, Any]) -> Paragraph:
    return Paragraph(escape(text_value), styles["section"])


def image_is_readable(path: str) -> bool:
    try:
        reader = ImageReader(path)
        reader.getSize()
        return True
    except Exception:
        return False


def build_standardized_resume_pdf(candidate: Dict[str, Any], company: Dict[str, Any], resume_documents: Optional[List[Dict[str, Any]]] = None, login_contact: Optional[Dict[str, Any]] = None) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=38,
        leftMargin=38,
        topMargin=34,
        bottomMargin=34,
        title=f"{candidate.get('candidate_code', 'Candidate')} Standard Resume",
    )

    sample = getSampleStyleSheet()
    styles = {
        "title": ParagraphStyle("Title", parent=sample["Title"], fontName="Helvetica-Bold", fontSize=18, leading=22, alignment=TA_LEFT, textColor=colors.HexColor("#0b0835")),
        "subtitle": ParagraphStyle("Subtitle", parent=sample["Normal"], fontSize=9.5, leading=12, textColor=colors.HexColor("#334155")),
        "section": ParagraphStyle("Section", parent=sample["Heading2"], fontName="Helvetica-Bold", fontSize=11.5, leading=14, textColor=colors.white, backColor=colors.HexColor("#0b0835"), spaceBefore=9, spaceAfter=6, leftIndent=0, borderPadding=5),
        "body": ParagraphStyle("Body", parent=sample["BodyText"], fontSize=8.7, leading=11, textColor=colors.HexColor("#0f172a")),
        "small": ParagraphStyle("Small", parent=sample["BodyText"], fontSize=8.2, leading=10, textColor=colors.HexColor("#0f172a")),
        "muted": ParagraphStyle("Muted", parent=sample["BodyText"], fontSize=7.5, leading=9, textColor=colors.HexColor("#64748b")),
        "name": ParagraphStyle("Name", parent=sample["Heading1"], fontName="Helvetica-Bold", fontSize=16, leading=20, textColor=colors.HexColor("#0f172a")),
    }

    story = []
    login_contact = login_contact or {}
    authorized_name = authorized_contact_value(login_contact.get("full_name"))
    authorized_email = authorized_contact_value(login_contact.get("email"))
    authorized_phone = authorized_contact_value(login_contact.get("phone"))
    company_name = clean_resume_value(company.get("company_name"), "Truflux Technologies")
    company_lines = [
        company_name,
        clean_resume_value(company.get("address"), ""),
        " | ".join([x for x in [clean_resume_value(company.get("phone"), ""), clean_resume_value(company.get("email"), ""), clean_resume_value(company.get("website"), "")] if x]),
        " | ".join([x for x in [f"Company No: {company.get('company_number')}" if company.get("company_number") else "", f"Tax/GST: {company.get('tax_number')}" if company.get("tax_number") else ""] if x]),
    ]
    company_para = "<br/>".join(escape(x) for x in company_lines if x)

    header_cells = []
    logo_path = company.get("logo_file_path")
    if logo_path and os.path.exists(logo_path) and image_is_readable(ensure_safe_file_path(logo_path)):
        try:
            header_cells.append(RLImage(ensure_safe_file_path(logo_path), width=0.9 * inch, height=0.55 * inch, kind="proportional"))
        except Exception:
            header_cells.append(Paragraph("<b>1Resource</b>", styles["subtitle"]))
    else:
        header_cells.append(Paragraph("<b>1Resource</b>", styles["subtitle"]))

    header_cells.append(Paragraph(f"<b>{escape(company_name)}</b><br/>{company_para}", styles["subtitle"]))
    header = Table([header_cells], colWidths=[1.35 * inch, 4.85 * inch])
    header.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f8fafc")),
        ("BOX", (0, 0), (-1, -1), 0.75, colors.HexColor("#dbe4f0")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(header)
    story.append(Spacer(1, 10))

    photo_cell = Paragraph("", styles["body"])
    photo_path = candidate.get("photo_file_path")
    if photo_path and os.path.exists(photo_path) and image_is_readable(ensure_safe_file_path(photo_path)):
        try:
            photo_cell = RLImage(ensure_safe_file_path(photo_path), width=0.9 * inch, height=1.05 * inch, kind="proportional")
        except Exception:
            photo_cell = Paragraph("Photo<br/>available", styles["muted"])
    candidate_title = Paragraph(
        f"<b>{pdf_text(candidate.get('full_name'), 'Candidate')}</b><br/>"
        f"{pdf_text(candidate.get('candidate_code'))} | {pdf_text(candidate.get('primary_skill'), 'Professional')}<br/>"
        f"Authorized contact: {pdf_text(authorized_email)} | {pdf_text(authorized_phone)}<br/>"
        f"Location: {pdf_text(candidate.get('location'))}",
        styles["name"],
    )
    profile_head = Table([[candidate_title, photo_cell]], colWidths=[5.0 * inch, 1.2 * inch])
    profile_head.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("BOX", (0, 0), (-1, -1), 0.75, colors.HexColor("#dbe4f0")),
        ("BACKGROUND", (0, 0), (-1, -1), colors.white),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(profile_head)

    story.append(section_title("Candidate Profile", styles))
    story.append(make_kv_table([
        ("Authorized Contact Name", authorized_name),
        ("Authorized Contact Email", authorized_email),
        ("Authorized Contact Phone", authorized_phone),
        ("Candidate Contact", CONTACT_SUPPRESSED),
        ("Location", candidate.get("location")),
        ("Availability", f"{clean_resume_value(candidate.get('current_status'))} | Available by: {clean_resume_value(candidate.get('available_by_date'))} | Notice: {clean_resume_value(candidate.get('notice_period_days'), '0')} days"),
        ("Employment Type", candidate.get("employment_type")),
        ("Recruiter Owner", candidate.get("recruiter_owner")),
    ], styles))

    story.append(section_title("Professional Summary", styles))
    story.append(make_kv_table([
        ("Primary Role / Skill", candidate.get("primary_skill")),
        ("Secondary Skills", candidate.get("secondary_skills")),
        ("Total Experience", f"{clean_resume_value(candidate.get('total_experience'), '0')} years"),
        ("Relevant Experience", f"{clean_resume_value(candidate.get('relevant_experience'), '0')} years"),
        ("Internal Level", candidate.get("internal_level")),
        ("Proficiency", candidate.get("proficiency")),
    ], styles))

    story.append(section_title("Company and Project Details", styles))
    story.append(make_kv_table([
        ("Current / Last Company", candidate.get("current_company")),
        ("Previous Companies", candidate.get("previous_companies")),
        ("Project / Assignment Details", redact_contact_info(candidate.get("project_details") or candidate.get("resume_text"))),
        ("Domain Exposure", candidate.get("domain_exposure")),
        ("Certifications", candidate.get("certifications")),
        ("Portfolio", candidate.get("portfolio_url")),
    ], styles))

    resume_documents = resume_documents or []
    if resume_documents:
        story.append(section_title("Uploaded Resume Document Details", styles))
        story.append(Paragraph("Candidate email IDs and phone numbers from uploaded resumes are suppressed. Use the authorized contact from the login profile shown in this PDF.", styles["muted"]))
        story.append(Spacer(1, 5))
        for idx, resume in enumerate(resume_documents[:5], start=1):
            analysis = parse_resume_analysis_json(resume.get("extracted_json"))
            detected_skills = ", ".join(analysis.get("skills") or []) if isinstance(analysis.get("skills"), list) else ""
            detected_domains = ", ".join(analysis.get("domains") or []) if isinstance(analysis.get("domains"), list) else ""
            detected_email = CONTACT_SUPPRESSED if (analysis.get("email") or "").strip() else "-"
            detected_phone = CONTACT_SUPPRESSED if (analysis.get("phone") or "").strip() else "-"
            detected_exp = analysis.get("experience_years")
            extracted_text = resume.get("resume_text") or ""
            rows = [
                ("Document", f"{idx}. {clean_resume_value(resume.get('resume_file_name'), 'Uploaded resume')}"),
                ("Role Title", resume.get("role_title")),
                ("Source / Uploaded On", f"{clean_resume_value(resume.get('source'))} | {clean_resume_value(resume.get('created_at'))}"),
                ("Detected Contact", f"Email: {clean_resume_value(detected_email)} | Phone: {clean_resume_value(detected_phone)}"),
                ("Detected Experience", f"{clean_resume_value(detected_exp, '0')} years"),
                ("Detected Skills", detected_skills or resume.get("skill_matches") or candidate.get("skill_matches")),
                ("Detected Domains", detected_domains or candidate.get("domain_exposure")),
                ("Fit / Risk", f"Fit {clean_resume_value(resume.get('fit_score'), '0')}/100 - {clean_resume_value(resume.get('rating_level'), 'Not rated')} | Risk {clean_resume_value(resume.get('fake_risk_score'), '0')}/100 - {clean_resume_value(resume.get('fake_risk_level'), 'Not checked')}"),
                ("Skill Matches", resume.get("skill_matches")),
                ("Skill Gaps", resume.get("skill_gaps")),
                ("Risk Reasons", resume.get("fake_risk_reasons")),
                ("Extracted Resume Content", compact_resume_text(extracted_text)),
            ]
            story.append(make_kv_table(rows, styles))
            story.append(Spacer(1, 7))
    else:
        story.append(section_title("Uploaded Resume Document Details", styles))
        story.append(Paragraph("No uploaded resume document text is available yet. Upload a resume version to enrich this section.", styles["muted"]))

    story.append(section_title("1Resource Screening Summary", styles))
    story.append(make_kv_table([
        ("Resume Fit Score", f"{clean_resume_value(candidate.get('ml_rating_score'), '0')}/100 - {clean_resume_value(candidate.get('ml_rating_level'), 'Not rated')}"),
        ("Fake Risk Score", f"{clean_resume_value(candidate.get('fake_risk_score'), '0')}/100 - {clean_resume_value(candidate.get('fake_risk_level'), 'Not checked')}"),
        ("Fake Risk RAG", clean_resume_value(candidate.get("fake_risk_rag"), "Not available")),
        ("Skill Matches", candidate.get("skill_matches")),
        ("Skill Gaps", candidate.get("skill_gaps")),
        ("Risk Reasons", candidate.get("fake_risk_reasons")),
    ], styles))

    story.append(section_title("Commercial Details", styles))
    story.append(make_kv_table([
        ("Expected Rate", candidate.get("expected_rate")),
        ("Negotiated Rate", candidate.get("negotiated_rate")),
    ], styles))

    story.append(Spacer(1, 8))
    story.append(Paragraph(
        f"Generated on {escape(now_iso())}. This resume is generated in the standardized 1Resource format by {escape(company_name)}. Authorized contact: {escape(authorized_email)} | {escape(authorized_phone)}.",
        styles["muted"],
    ))

    doc.build(story)
    return buffer.getvalue()


@app.get("/api/health")
def health() -> Dict[str, str]:
    return {
        "status": "ok",
        "app": APP_NAME,
        "version": APP_VERSION,
        "app_env": APP_ENV,
        "database": "postgresql" if USE_POSTGRES else "sqlite",
        "data_dir": DATA_DIR,
    }


@app.post("/api/login")
def login(payload: LoginRequest, request: Request) -> Dict[str, Any]:
    username = payload.username.strip()
    with get_db() as conn:
        row = conn.execute("SELECT * FROM users WHERE username = ?", (username,)).fetchone()
        if row and row["locked_until"]:
            locked_until = parse_optional_iso(row["locked_until"])
            if locked_until and locked_until > datetime.utcnow():
                raise HTTPException(status_code=423, detail="Account is temporarily locked after repeated failed attempts. Try again later or ask an admin to unlock it.")
        if not row or row["is_active"] != 1 or not verify_password(payload.password, row["password_hash"]):
            if row:
                attempts = int(row["failed_attempts"] or 0) + 1
                locked_until = (datetime.utcnow() + timedelta(minutes=15)).replace(microsecond=0).isoformat() + "Z" if attempts >= 5 else None
                conn.execute("UPDATE users SET failed_attempts=?, locked_until=? WHERE id=?", (attempts, locked_until, row["id"]))
                conn.commit()
                log_action(username, "failed_login", "user", row["id"], f"attempts={attempts}")
            raise HTTPException(status_code=401, detail="Invalid username or password")
        token = secrets.token_urlsafe(36)
        expires_at = session_expiry()
        conn.execute(
            "INSERT INTO sessions(token, user_id, created_at, expires_at, last_seen_at, user_agent, ip_address) VALUES(?,?,?,?,?,?,?)",
            (token, row["id"], now_iso(), expires_at, now_iso(), request.headers.get("user-agent", "")[:300], get_client_ip(request)),
        )
        conn.execute("UPDATE users SET failed_attempts=0, locked_until=NULL, last_login_at=? WHERE id=?", (now_iso(), row["id"]))
        conn.commit()
    log_action(username, "login", "user", row["id"], f"expires={expires_at}")
    return {"token": token, "expires_at": expires_at, "user": {"id": row["id"], "username": row["username"], "full_name": row["full_name"], "email": row["email"] or "", "phone": row["phone"] or "", "role": row["role"], "force_password_change": row["force_password_change"]}}


@app.post("/api/logout")
def logout(authorization: str = Header(default=""), user: Dict[str, Any] = Depends(get_current_user)) -> Dict[str, str]:
    token = authorization.replace("Bearer ", "", 1).strip()
    with get_db() as conn:
        conn.execute("DELETE FROM sessions WHERE token = ?", (token,))
        conn.commit()
    log_action(user["username"], "logout", "user", user["id"])
    return {"message": "Logged out"}


@app.get("/api/me")
def me(user: Dict[str, Any] = Depends(get_current_user)) -> Dict[str, Any]:
    data = dict(user)
    data.pop("photo_file_path", None)
    return data


@app.put("/api/me/profile")
def update_my_profile(payload: UserProfileUpdate, user: Dict[str, Any] = Depends(get_current_user)) -> Dict[str, Any]:
    full_name = (payload.full_name or user.get("full_name") or "").strip()
    email = (payload.email or "").strip()
    phone = (payload.phone or "").strip()
    if not full_name:
        raise HTTPException(status_code=400, detail="Full name is required")
    with get_db() as conn:
        conn.execute(
            "UPDATE users SET full_name=?, email=?, phone=? WHERE id=?",
            (full_name, email, phone, user["id"]),
        )
        conn.commit()
        row = conn.execute("SELECT id, username, full_name, email, phone, photo_file_name, photo_file_path, role, is_active, force_password_change FROM users WHERE id=?", (user["id"],)).fetchone()
    log_action(user["username"], "update_login_profile", "user", user["id"])
    data = dict_row(row)
    data.pop("photo_file_path", None)
    return data


@app.post("/api/me/photo")
def upload_my_photo(photo: UploadFile = File(...), user: Dict[str, Any] = Depends(get_current_user)) -> Dict[str, Any]:
    photo_name, photo_path = save_photo_upload(photo, f"user_{user['id']}_photo")
    with get_db() as conn:
        conn.execute("UPDATE users SET photo_file_name=?, photo_file_path=? WHERE id=?", (photo_name, photo_path, user["id"]))
        conn.commit()
        row = conn.execute(
            "SELECT id, username, full_name, email, phone, photo_file_name, photo_file_path, role, is_active, force_password_change FROM users WHERE id=?",
            (user["id"],),
        ).fetchone()
    log_action(user["username"], "upload_user_photo", "user", user["id"], photo_name)
    data = dict_row(row)
    data.pop("photo_file_path", None)
    return data


@app.get("/api/me/photo")
def get_my_photo(user: Dict[str, Any] = Depends(get_current_user)):
    path = user.get("photo_file_path") or ""
    if not path or not os.path.exists(path):
        raise HTTPException(status_code=404, detail="Profile photo not found")
    return FileResponse(path, filename=user.get("photo_file_name") or "profile_photo")


class PasswordChange(BaseModel):
    current_password: str
    new_password: str


class PasswordReset(BaseModel):
    new_password: str
    force_change: bool = True


@app.post("/api/change-password")
def change_password(payload: PasswordChange, user: Dict[str, Any] = Depends(get_current_user)) -> Dict[str, str]:
    enforce_password_policy(payload.new_password)
    with get_db() as conn:
        row = conn.execute("SELECT * FROM users WHERE id=?", (user["id"],)).fetchone()
        if not row or not verify_password(payload.current_password, row["password_hash"]):
            raise HTTPException(status_code=401, detail="Current password is incorrect")
        conn.execute("UPDATE users SET password_hash=?, force_password_change=0, password_changed_at=?, failed_attempts=0, locked_until=NULL WHERE id=?", (hash_password(payload.new_password), now_iso(), user["id"]))
        conn.execute("DELETE FROM sessions WHERE user_id=? AND token NOT IN (SELECT token FROM sessions WHERE user_id=? ORDER BY created_at DESC LIMIT 1)", (user["id"], user["id"]))
        conn.commit()
    log_action(user["username"], "change_password", "user", user["id"])
    return {"message": "Password changed successfully"}


@app.get("/api/company-profile")
def get_company_profile(user: Dict[str, Any] = Depends(get_current_user)) -> Dict[str, Any]:
    with get_db() as conn:
        row = conn.execute("SELECT * FROM company_profile WHERE id=1").fetchone()
        if not row:
            conn.execute(
                """
                INSERT OR IGNORE INTO company_profile(id, company_name, company_number, tax_number, address, phone, email, website, updated_at)
                VALUES(1, 'Truflux Technologies', '', '', '', '', '', '', ?)
                """,
                (now_iso(),),
            )
            conn.commit()
            row = conn.execute("SELECT * FROM company_profile WHERE id=1").fetchone()
    return dict_row(row)


@app.put("/api/company-profile")
def update_company_profile(payload: CompanyProfileUpdate, user: Dict[str, Any] = Depends(require_roles("Admin"))) -> Dict[str, Any]:
    now = now_iso()
    clean_payload = {
        "company_name": (payload.company_name or "Truflux Technologies").strip(),
        "company_number": (payload.company_number or "").strip(),
        "tax_number": (payload.tax_number or "").strip(),
        "address": (payload.address or "").strip(),
        "phone": (payload.phone or "").strip(),
        "email": (payload.email or "").strip(),
        "website": (payload.website or "").strip(),
    }
    with get_db() as conn:
        existing = conn.execute("SELECT id FROM company_profile WHERE id=1").fetchone()
        if not existing:
            conn.execute(
                """
                INSERT INTO company_profile(id, company_name, company_number, tax_number, address, phone, email, website, updated_at)
                VALUES(1,?,?,?,?,?,?,?,?)
                """,
                (
                    clean_payload["company_name"],
                    clean_payload["company_number"],
                    clean_payload["tax_number"],
                    clean_payload["address"],
                    clean_payload["phone"],
                    clean_payload["email"],
                    clean_payload["website"],
                    now,
                ),
            )
        else:
            conn.execute(
                """
                UPDATE company_profile
                SET company_name=?, company_number=?, tax_number=?, address=?, phone=?, email=?, website=?, updated_at=?
                WHERE id=1
                """,
                (
                    clean_payload["company_name"],
                    clean_payload["company_number"],
                    clean_payload["tax_number"],
                    clean_payload["address"],
                    clean_payload["phone"],
                    clean_payload["email"],
                    clean_payload["website"],
                    now,
                ),
            )
        conn.commit()
        row = conn.execute("SELECT * FROM company_profile WHERE id=1").fetchone()
    log_action(user["username"], "update_company_profile", "company_profile", 1, json.dumps(clean_payload))
    return dict_row(row)




@app.post("/api/company-profile/logo")
def upload_company_logo(logo: UploadFile = File(...), user: Dict[str, Any] = Depends(require_roles("Admin"))) -> Dict[str, Any]:
    logo_name, logo_path = save_photo_upload(logo, "company_logo")
    with get_db() as conn:
        conn.execute(
            """
            INSERT OR IGNORE INTO company_profile(id, company_name, updated_at)
            VALUES(1, 'Truflux Technologies', ?)
            """,
            (now_iso(),),
        )
        conn.execute(
            """
            UPDATE company_profile SET logo_file_name=?, logo_file_path=?, updated_at=? WHERE id=1
            """,
            (logo_name, logo_path, now_iso()),
        )
        conn.commit()
    log_action(user["username"], "upload_company_logo", "company_profile", 1, logo_name)
    return get_company_profile(user)


@app.get("/api/users")
def list_users(user: Dict[str, Any] = Depends(require_roles("Admin"))) -> List[Dict[str, Any]]:
    with get_db() as conn:
        rows = conn.execute("SELECT id, username, full_name, email, phone, role, is_active, failed_attempts, locked_until, force_password_change, last_login_at, created_at FROM users ORDER BY id DESC").fetchall()
    return [enrich_fake_risk_rag(dict_row(r)) for r in rows]


@app.post("/api/users")
def create_user(payload: UserCreate, user: Dict[str, Any] = Depends(require_roles("Admin"))) -> Dict[str, Any]:
    if payload.role not in ["Admin", "Recruiter", "Evaluator", "Viewer"]:
        raise HTTPException(status_code=400, detail="Invalid role")
    enforce_password_policy(payload.password)
    try:
        with get_db() as conn:
            cur = conn.execute(
                "INSERT INTO users(username, full_name, email, phone, role, password_hash, is_active, force_password_change, password_changed_at, created_at) VALUES(?,?,?,?,?,?,?,?,?,?)",
                (payload.username.strip(), payload.full_name.strip(), (payload.email or "").strip(), (payload.phone or "").strip(), payload.role, hash_password(payload.password), 1, 1, now_iso(), now_iso()),
            )
            conn.commit()
            user_id = cur.lastrowid
    except sqlite3.IntegrityError:
        raise HTTPException(status_code=400, detail="Username already exists")
    log_action(user["username"], "create_user", "user", user_id)
    return {"id": user_id, "username": payload.username.strip(), "full_name": payload.full_name.strip(), "email": (payload.email or "").strip(), "phone": (payload.phone or "").strip(), "role": payload.role, "is_active": 1}


@app.put("/api/users/{user_id}")
def update_user_details(user_id: int, payload: UserUpdate, user: Dict[str, Any] = Depends(require_roles("Admin"))) -> Dict[str, Any]:
    username = (payload.username or "").strip()
    full_name = (payload.full_name or "").strip()
    email = (payload.email or "").strip()
    phone = (payload.phone or "").strip()
    role = (payload.role or "Recruiter").strip()
    if not username:
        raise HTTPException(status_code=400, detail="Username is required")
    if not full_name:
        raise HTTPException(status_code=400, detail="Full name is required")
    if role not in ["Admin", "Recruiter", "Evaluator", "Viewer"]:
        raise HTTPException(status_code=400, detail="Invalid role")
    try:
        with get_db() as conn:
            existing = conn.execute("SELECT id FROM users WHERE id=?", (user_id,)).fetchone()
            if not existing:
                raise HTTPException(status_code=404, detail="User not found")
            duplicate = conn.execute("SELECT id FROM users WHERE username=? AND id<>?", (username, user_id)).fetchone()
            if duplicate:
                raise HTTPException(status_code=400, detail="Username already exists")
            conn.execute(
                """
                UPDATE users
                SET username=?, full_name=?, email=?, phone=?, role=?
                WHERE id=?
                """,
                (username, full_name, email, phone, role, user_id),
            )
            conn.commit()
            row = conn.execute(
                "SELECT id, username, full_name, email, phone, role, is_active, failed_attempts, locked_until, force_password_change, last_login_at, created_at FROM users WHERE id=?",
                (user_id,),
            ).fetchone()
    except sqlite3.IntegrityError:
        raise HTTPException(status_code=400, detail="Username already exists")
    log_action(user["username"], "update_user_details", "user", user_id, f"username={username}; role={role}")
    return dict_row(row)


@app.patch("/api/users/{user_id}/status")
def update_user_status(user_id: int, is_active: int, user: Dict[str, Any] = Depends(require_roles("Admin"))) -> Dict[str, str]:
    with get_db() as conn:
        conn.execute("UPDATE users SET is_active = ? WHERE id = ?", (1 if is_active else 0, user_id))
        conn.commit()
    log_action(user["username"], "update_user_status", "user", user_id, f"is_active={is_active}")
    return {"message": "User status updated"}


@app.post("/api/users/{user_id}/reset-password")
def reset_user_password(user_id: int, payload: PasswordReset, user: Dict[str, Any] = Depends(require_roles("Admin"))) -> Dict[str, str]:
    enforce_password_policy(payload.new_password)
    with get_db() as conn:
        exists = conn.execute("SELECT id FROM users WHERE id=?", (user_id,)).fetchone()
        if not exists:
            raise HTTPException(status_code=404, detail="User not found")
        conn.execute("UPDATE users SET password_hash=?, force_password_change=?, password_changed_at=?, failed_attempts=0, locked_until=NULL WHERE id=?", (hash_password(payload.new_password), 1 if payload.force_change else 0, now_iso(), user_id))
        conn.execute("DELETE FROM sessions WHERE user_id=?", (user_id,))
        conn.commit()
    log_action(user["username"], "reset_user_password", "user", user_id)
    return {"message": "Password reset successfully and active sessions cleared"}


@app.post("/api/users/{user_id}/unlock")
def unlock_user(user_id: int, user: Dict[str, Any] = Depends(require_roles("Admin"))) -> Dict[str, str]:
    with get_db() as conn:
        conn.execute("UPDATE users SET failed_attempts=0, locked_until=NULL WHERE id=?", (user_id,))
        conn.commit()
    log_action(user["username"], "unlock_user", "user", user_id)
    return {"message": "User unlocked"}


@app.get("/api/dashboard")
def dashboard(user: Dict[str, Any] = Depends(get_current_user)) -> Dict[str, Any]:
    with get_db() as conn:
        total = conn.execute("SELECT COUNT(*) AS c FROM candidates").fetchone()["c"]
        ready = conn.execute("SELECT COUNT(*) AS c FROM candidates WHERE status LIKE 'A1%' OR status LIKE 'A2%'").fetchone()["c"]
        bench = conn.execute("SELECT COUNT(*) AS c FROM candidates WHERE current_status IN ('Available','Bench','Freelance')").fetchone()["c"]
        avg_score = conn.execute("SELECT COALESCE(ROUND(AVG(total_score),1),0) AS s FROM assessments").fetchone()["s"]
        avg_ml = conn.execute("SELECT COALESCE(ROUND(AVG(ml_rating_score),1),0) AS s FROM candidates WHERE ml_rating_score > 0").fetchone()["s"]
        avg_risk = conn.execute("SELECT COALESCE(ROUND(AVG(fake_risk_score),1),0) AS s FROM candidates WHERE fake_risk_score > 0").fetchone()["s"]
        total_demand = conn.execute("SELECT COUNT(*) AS c FROM demand_requests").fetchone()["c"]
        open_demand = conn.execute("SELECT COUNT(*) AS c FROM demand_requests WHERE status IN ('Open','Hot','In Progress')").fetchone()["c"]
        shortlisted = conn.execute("SELECT COUNT(*) AS c FROM demand_shortlists").fetchone()["c"]
        skill_rows = conn.execute("SELECT COALESCE(primary_skill,'Unspecified') AS skill, COUNT(*) AS count FROM candidates GROUP BY primary_skill ORDER BY count DESC LIMIT 8").fetchall()
        demand_skill_rows = conn.execute("SELECT role_title, required_skills, role_definition FROM demand_requests WHERE status IN ('Open','Hot','In Progress') ORDER BY updated_at DESC LIMIT 50").fetchall()
        demand_counter = Counter()
        for r in demand_skill_rows:
            for sk in extract_skills(f"{r['role_title'] or ''} {r['required_skills'] or ''} {r['role_definition'] or ''}"):
                demand_counter[sk] += 1
        status_rows = conn.execute("SELECT status, COUNT(*) AS count FROM candidates GROUP BY status ORDER BY count DESC").fetchall()
        recent = conn.execute("SELECT id, candidate_code, full_name, primary_skill, status, updated_at, ml_rating_score, fake_risk_level FROM candidates ORDER BY updated_at DESC LIMIT 8").fetchall()
        recent_demand = conn.execute("SELECT id, demand_code, client_name, project_name, role_title, status, priority, updated_at FROM demand_requests ORDER BY updated_at DESC LIMIT 8").fetchall()
    return {
        "total_candidates": total,
        "ready_candidates": ready,
        "available_bench": bench,
        "average_score": avg_score,
        "average_ml_rating": avg_ml,
        "average_fake_risk": avg_risk,
        "total_demand": total_demand,
        "open_demand": open_demand,
        "shortlisted": shortlisted,
        "skills": [dict_row(r) for r in skill_rows],
        "demand_skills": [{"skill": k, "count": v} for k, v in demand_counter.most_common(8)],
        "statuses": [dict_row(r) for r in status_rows],
        "recent_candidates": [enrich_fake_risk_rag(dict_row(r)) for r in recent],
        "recent_demand": [dict_row(r) for r in recent_demand],
    }


@app.get("/api/candidates")
def list_candidates(q: str = "", skill: str = "", status: str = "", availability: str = "", limit: int = 100, user: Dict[str, Any] = Depends(get_current_user)) -> List[Dict[str, Any]]:
    clauses = []
    values: List[Any] = []
    if q:
        clauses.append("(full_name LIKE ? OR email LIKE ? OR candidate_code LIKE ? OR secondary_skills LIKE ? OR domain_exposure LIKE ?)")
        like = f"%{q}%"
        values.extend([like, like, like, like, like])
    if skill:
        clauses.append("(primary_skill LIKE ? OR secondary_skills LIKE ? OR skill_matches LIKE ?)")
        like = f"%{skill}%"
        values.extend([like, like, like])
    if status:
        clauses.append("status = ?")
        values.append(status)
    if availability:
        clauses.append("current_status = ?")
        values.append(availability)
    where = " WHERE " + " AND ".join(clauses) if clauses else ""
    values.append(limit)
    with get_db() as conn:
        rows = conn.execute(f"SELECT * FROM candidates {where} ORDER BY updated_at DESC LIMIT ?", values).fetchall()
    return [enrich_fake_risk_rag(dict_row(r)) for r in rows]


@app.get("/api/candidates/{candidate_id}")
def get_candidate(candidate_id: int, user: Dict[str, Any] = Depends(get_current_user)) -> Dict[str, Any]:
    with get_db() as conn:
        c = conn.execute("SELECT * FROM candidates WHERE id = ?", (candidate_id,)).fetchone()
        if not c:
            raise HTTPException(status_code=404, detail="Candidate not found")
        assessments = conn.execute("SELECT * FROM assessments WHERE candidate_id = ? ORDER BY created_at DESC", (candidate_id,)).fetchall()
        resumes = conn.execute("SELECT id, role_title, role_definition, resume_file_name, fit_score, rating_level, fake_risk_score, fake_risk_level, fake_risk_reasons, skill_matches, skill_gaps, source, created_at FROM resumes WHERE candidate_id = ? ORDER BY created_at DESC", (candidate_id,)).fetchall()
        mcq_results = conn.execute(
            """
            SELECT pm.*, d.demand_code, d.client_name, d.project_name, d.role_title
            FROM public_mcq_results pm
            LEFT JOIN demand_requests d ON pm.demand_id = d.id
            WHERE pm.candidate_id = ?
            ORDER BY pm.created_at DESC
            """,
            (candidate_id,),
        ).fetchall()
    out = enrich_fake_risk_rag(dict_row(c))
    out["assessments"] = [dict_row(a) for a in assessments]
    out["resumes"] = [dict_row(r) for r in resumes]
    out["mcq_results"] = [dict_row(m) for m in mcq_results]
    return out


@app.post("/api/candidates")
def create_candidate(payload: CandidateCreate, user: Dict[str, Any] = Depends(require_roles("Admin", "Recruiter"))) -> Dict[str, Any]:
    with get_db() as conn:
        code = generate_candidate_code(conn)
        cur = conn.execute(
            """
            INSERT INTO candidates(candidate_code, full_name, email, phone, location, current_status, availability_date, available_by_date, notice_period_days,
                employment_type, source, recruiter_owner, total_experience, relevant_experience, primary_skill,
                secondary_skills, domain_exposure, proficiency, certifications, portfolio_url, current_company, previous_companies, project_details, expected_rate, negotiated_rate,
                internal_level, resume_text, status, created_date, created_at, updated_at)
            VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
            """,
            (code, payload.full_name, payload.email, payload.phone, payload.location, payload.current_status, payload.availability_date, payload.available_by_date, int(payload.notice_period_days or 0),
             payload.employment_type, payload.source, payload.recruiter_owner or user["full_name"], payload.total_experience,
             payload.relevant_experience, payload.primary_skill, payload.secondary_skills, payload.domain_exposure, payload.proficiency,
             payload.certifications, payload.portfolio_url, payload.current_company, payload.previous_companies, payload.project_details,
             payload.expected_rate, payload.negotiated_rate, payload.internal_level,
             payload.resume_text, payload.status, today_date(), now_iso(), now_iso()),
        )
        conn.commit()
        candidate_id = cur.lastrowid
    log_action(user["username"], "create_candidate", "candidate", candidate_id, code)
    return get_candidate(candidate_id, user)


@app.put("/api/candidates/{candidate_id}")
def update_candidate(candidate_id: int, payload: CandidateUpdate, user: Dict[str, Any] = Depends(require_roles("Admin", "Recruiter"))) -> Dict[str, Any]:
    with get_db() as conn:
        existing = conn.execute("SELECT id FROM candidates WHERE id = ?", (candidate_id,)).fetchone()
        if not existing:
            raise HTTPException(status_code=404, detail="Candidate not found")
        conn.execute(
            """
            UPDATE candidates SET full_name=?, email=?, phone=?, location=?, current_status=?, availability_date=?, available_by_date=?, notice_period_days=?, employment_type=?,
                source=?, recruiter_owner=?, total_experience=?, relevant_experience=?, primary_skill=?, secondary_skills=?,
                domain_exposure=?, proficiency=?, certifications=?, portfolio_url=?, current_company=?, previous_companies=?, project_details=?, expected_rate=?, negotiated_rate=?,
                internal_level=?, resume_text=?, status=?, updated_at=? WHERE id=?
            """,
            (payload.full_name, payload.email, payload.phone, payload.location, payload.current_status, payload.availability_date, payload.available_by_date, int(payload.notice_period_days or 0),
             payload.employment_type, payload.source, payload.recruiter_owner, payload.total_experience, payload.relevant_experience,
             payload.primary_skill, payload.secondary_skills, payload.domain_exposure, payload.proficiency, payload.certifications,
             payload.portfolio_url, payload.current_company, payload.previous_companies, payload.project_details,
             payload.expected_rate, payload.negotiated_rate, payload.internal_level, payload.resume_text,
             payload.status, now_iso(), candidate_id),
        )
        conn.commit()
    log_action(user["username"], "update_candidate", "candidate", candidate_id)
    return get_candidate(candidate_id, user)


@app.delete("/api/candidates/{candidate_id}")
def delete_candidate(candidate_id: int, user: Dict[str, Any] = Depends(require_roles("Admin"))) -> Dict[str, str]:
    with get_db() as conn:
        conn.execute("DELETE FROM candidates WHERE id = ?", (candidate_id,))
        conn.commit()
    log_action(user["username"], "delete_candidate", "candidate", candidate_id)
    return {"message": "Candidate deleted"}


@app.post("/api/candidates/{candidate_id}/resumes")
def upload_role_resume(candidate_id: int, role_title: str = Form(""), role_definition: str = Form(""), file: UploadFile = File(...), user: Dict[str, Any] = Depends(require_roles("Admin", "Recruiter"))) -> Dict[str, Any]:
    file_name, file_path = save_upload(file, f"candidate_{candidate_id}")
    resume_text = extract_text_from_file(file_path, file_name)
    analysis = analyze_resume_text(resume_text, role_definition, role_title)
    with get_db() as conn:
        existing = conn.execute("SELECT id FROM candidates WHERE id = ?", (candidate_id,)).fetchone()
        if not existing:
            raise HTTPException(status_code=404, detail="Candidate not found")
        resume_id = create_resume_record(conn, candidate_id, file_name, file_path, resume_text, role_title, role_definition, "Recruiter Upload", analysis)
        conn.commit()
    log_action(user["username"], "upload_role_resume", "candidate", candidate_id, f"resume_id={resume_id}; fit={analysis['fit_score']}; risk={analysis['fake_risk_score']}")
    return {"message": "Resume uploaded and candidate profile updated", "resume_id": resume_id, "analysis": analysis}


@app.post("/api/candidates/{candidate_id}/upload")
def upload_resume(candidate_id: int, file: UploadFile = File(...), user: Dict[str, Any] = Depends(require_roles("Admin", "Recruiter"))) -> Dict[str, str]:
    # Backward-compatible endpoint from older versions.
    file_name, file_path = save_upload(file, f"candidate_{candidate_id}")
    resume_text = extract_text_from_file(file_path, file_name)
    analysis = analyze_resume_text(resume_text, "", "")
    with get_db() as conn:
        create_resume_record(conn, candidate_id, file_name, file_path, resume_text, "General Resume", "", "Recruiter Upload", analysis)
        conn.commit()
    log_action(user["username"], "upload_resume", "candidate", candidate_id, file_name)
    return {"message": "Resume uploaded", "file_name": file_name}


@app.get("/api/candidates/{candidate_id}/download")
def download_resume(candidate_id: int, user: Dict[str, Any] = Depends(get_current_user)):
    with get_db() as conn:
        row = conn.execute("SELECT resume_file_name, resume_file_path FROM candidates WHERE id = ?", (candidate_id,)).fetchone()
    if not row or not row["resume_file_path"]:
        raise HTTPException(status_code=404, detail="Resume file not found")
    return FileResponse(path=ensure_safe_file_path(row["resume_file_path"]), filename=row["resume_file_name"])


@app.get("/api/resumes/{resume_id}/download")
def download_resume_version(resume_id: int, user: Dict[str, Any] = Depends(get_current_user)):
    with get_db() as conn:
        row = conn.execute("SELECT resume_file_name, resume_file_path FROM resumes WHERE id = ?", (resume_id,)).fetchone()
    if not row or not row["resume_file_path"]:
        raise HTTPException(status_code=404, detail="Resume file not found")
    return FileResponse(path=ensure_safe_file_path(row["resume_file_path"]), filename=row["resume_file_name"])


@app.get("/api/candidates/{candidate_id}/photo")
def download_candidate_photo(candidate_id: int, user: Dict[str, Any] = Depends(get_current_user)):
    with get_db() as conn:
        row = conn.execute("SELECT photo_file_name, photo_file_path FROM candidates WHERE id = ?", (candidate_id,)).fetchone()
    if not row or not row["photo_file_path"]:
        raise HTTPException(status_code=404, detail="Candidate photograph not found")
    return FileResponse(path=ensure_safe_file_path(row["photo_file_path"]), filename=row["photo_file_name"] or "candidate_photograph")


@app.get("/api/candidates/{candidate_id}/standard-resume")
def download_standard_resume(candidate_id: int, user: Dict[str, Any] = Depends(get_current_user)):
    with get_db() as conn:
        row = conn.execute("SELECT * FROM candidates WHERE id = ?", (candidate_id,)).fetchone()
        company_row = conn.execute("SELECT * FROM company_profile WHERE id=1").fetchone()
        resume_rows = conn.execute(
            """
            SELECT id, role_title, role_definition, resume_file_name, resume_text, extracted_json,
                   fit_score, rating_level, fake_risk_score, fake_risk_level, fake_risk_reasons,
                   skill_matches, skill_gaps, source, created_at
            FROM resumes
            WHERE candidate_id=?
            ORDER BY created_at DESC, id DESC
            """,
            (candidate_id,),
        ).fetchall()
    if not row:
        raise HTTPException(status_code=404, detail="Candidate not found")
    candidate = enrich_fake_risk_rag(dict_row(row))
    company = dict_row(company_row) if company_row else {"company_name": "Truflux Technologies"}
    resume_documents = [dict_row(r) for r in resume_rows]
    pdf_bytes = build_standardized_resume_pdf(candidate, company, resume_documents, login_contact=user)
    file_name = f"{sanitize_filename(candidate.get('candidate_code') or 'candidate')}_standard_resume.pdf"
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{file_name}"'}
    )


@app.post("/api/candidates/{candidate_id}/assessments")
def add_assessment(candidate_id: int, payload: AssessmentCreate, user: Dict[str, Any] = Depends(require_roles("Admin", "Evaluator"))) -> Dict[str, Any]:
    total = compute_total(payload)
    derived_status = status_from_score(total)
    with get_db() as conn:
        c = conn.execute("SELECT id FROM candidates WHERE id = ?", (candidate_id,)).fetchone()
        if not c:
            raise HTTPException(status_code=404, detail="Candidate not found")
        cur = conn.execute(
            """
            INSERT INTO assessments(candidate_id, technical_score, project_score, practical_score, communication_score,
                client_readiness_score, cost_fitment_score, availability_score, total_score, evaluator_name,
                recommendation, remarks, created_at)
            VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)
            """,
            (candidate_id, payload.technical_score, payload.project_score, payload.practical_score, payload.communication_score,
             payload.client_readiness_score, payload.cost_fitment_score, payload.availability_score, total,
             payload.evaluator_name or user["full_name"], payload.recommendation or derived_status, payload.remarks, now_iso()),
        )
        conn.execute("UPDATE candidates SET status=?, updated_at=? WHERE id=?", (derived_status, now_iso(), candidate_id))
        conn.commit()
        assessment_id = cur.lastrowid
        row = conn.execute("SELECT * FROM assessments WHERE id = ?", (assessment_id,)).fetchone()
    log_action(user["username"], "add_assessment", "candidate", candidate_id, f"score={total}")
    return dict_row(row)




def build_mcq_questions_for_demand(demand: Dict[str, Any]) -> List[Dict[str, Any]]:
    role = demand.get("role_title") or "Technology Role"
    combined = f"{role} {demand.get('required_skills') or ''} {demand.get('role_definition') or ''}"
    skills = extract_skills(combined)
    if not skills:
        # Add a few practical fallback topics so every demand gets 10 questions.
        skills = ["API", "SQL", "Git", "Security", "Testing"]
    while len(skills) < 10:
        skills = skills + skills
    templates = [
        ("What is the best first step when starting a {skill} task for {role}?", ["Clarify the requirement and acceptance criteria", "Start coding immediately", "Skip documentation", "Deploy directly to production"], "A", "Good delivery starts with clear requirements and acceptance criteria."),
        ("Which practice is most important for production-ready {skill} work?", ["Peer review and testing before release", "Changing code directly on the server", "Ignoring logs", "Using only manual memory checks"], "A", "Review and testing reduce delivery risk."),
        ("A client reports a defect in a {skill} feature. What should the resource do first?", ["Reproduce the issue and capture evidence", "Blame another team", "Delete recent changes", "Close the ticket without validation"], "A", "Reproduction and evidence help isolate the root cause."),
        ("For {skill}, what indicates stronger hands-on experience?", ["Explaining trade-offs from a real project", "Only listing buzzwords", "Avoiding architecture questions", "No examples of delivery"], "A", "Real experience shows through project examples and trade-offs."),
        ("What is a good way to estimate a {skill} deliverable?", ["Break it into tasks, assumptions, risks and dependencies", "Give a random number", "Estimate without understanding scope", "Ignore testing effort"], "A", "Task breakdown and assumptions make estimates defensible."),
        ("Which behavior is a red flag in a {skill} implementation?", ["No validation, no logs, and no error handling", "Unit tests for critical logic", "Code review", "Clear deployment checklist"], "A", "Missing validation, logs and error handling raises risk."),
        ("For {role}, how should {skill} progress be reported?", ["Clear status, blockers, risks, and next steps", "Only say everything is fine", "Avoid mentioning blockers", "Report after the deadline only"], "A", "Professional reporting should be transparent and actionable."),
        ("When using {skill}, why is version control important?", ["It tracks changes, supports rollback, and improves collaboration", "It replaces testing", "It removes the need for documentation", "It guarantees no defects"], "A", "Version control protects delivery traceability and collaboration."),
        ("Which answer best shows client-readiness in {skill}?", ["Explain the approach in business-friendly language", "Use only jargon", "Avoid questions", "Say the client will not understand"], "A", "Client-facing resources must explain clearly without unnecessary jargon."),
        ("What should be checked before marking a {skill} task complete?", ["Acceptance criteria, tests, edge cases, and deployment impact", "Only whether the screen opens", "Nothing after coding", "Only the developer's opinion"], "A", "Completion should be validated against agreed acceptance criteria."),
    ]
    questions = []
    for idx in range(10):
        skill = skills[idx % len(skills)]
        q, options, correct, explanation = templates[idx]
        # Rotate options in a deterministic way so correct answers are not visually always first.
        rotation = idx % 4
        rotated = options[rotation:] + options[:rotation]
        correct_text = options[0]
        correct_key = "ABCD"[rotated.index(correct_text)]
        questions.append({
            "question_no": idx + 1,
            "skill": skill,
            "question_text": q.format(skill=skill, role=role),
            "option_a": rotated[0],
            "option_b": rotated[1],
            "option_c": rotated[2],
            "option_d": rotated[3],
            "correct_option": correct_key,
            "explanation": explanation,
        })
    return questions


def get_or_create_demand_mcqs(conn: sqlite3.Connection, demand_id: int, refresh: bool = False) -> List[sqlite3.Row]:
    if refresh:
        conn.execute("DELETE FROM demand_mcq_questions WHERE demand_id=?", (demand_id,))
    rows = conn.execute("SELECT * FROM demand_mcq_questions WHERE demand_id=? ORDER BY question_no", (demand_id,)).fetchall()
    if len(rows) >= 10:
        return rows
    demand = conn.execute("SELECT * FROM demand_requests WHERE id=?", (demand_id,)).fetchone()
    if not demand:
        raise HTTPException(status_code=404, detail="Demand request not found")
    conn.execute("DELETE FROM demand_mcq_questions WHERE demand_id=?", (demand_id,))
    now = now_iso()
    for q in build_mcq_questions_for_demand(dict_row(demand)):
        conn.execute(
            """
            INSERT INTO demand_mcq_questions(demand_id, question_no, skill, question_text, option_a, option_b, option_c, option_d,
                correct_option, explanation, created_at, updated_at)
            VALUES(?,?,?,?,?,?,?,?,?,?,?,?)
            """,
            (demand_id, q["question_no"], q["skill"], q["question_text"], q["option_a"], q["option_b"], q["option_c"], q["option_d"], q["correct_option"], q["explanation"], now, now),
        )
    return conn.execute("SELECT * FROM demand_mcq_questions WHERE demand_id=? ORDER BY question_no", (demand_id,)).fetchall()


def safe_mcq_rows(rows: List[sqlite3.Row], include_answers: bool = False) -> List[Dict[str, Any]]:
    out = []
    for r in rows:
        item = {
            "id": r["id"],
            "question_no": r["question_no"],
            "skill": r["skill"],
            "question_text": r["question_text"],
            "options": [
                {"key": "A", "text": r["option_a"]},
                {"key": "B", "text": r["option_b"]},
                {"key": "C", "text": r["option_c"]},
                {"key": "D", "text": r["option_d"]},
            ],
            "explanation": r["explanation"],
        }
        if include_answers:
            item["correct_option"] = r["correct_option"]
        out.append(item)
    return out


def score_mcq_answers(conn: sqlite3.Connection, demand_id: int, answers_json: str) -> Dict[str, Any]:
    rows = get_or_create_demand_mcqs(conn, demand_id)
    try:
        raw = json.loads(answers_json or "{}")
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid MCQ answers payload")
    answers = raw.get("answers", raw) if isinstance(raw, dict) else {}
    if not isinstance(answers, dict):
        raise HTTPException(status_code=400, detail="MCQ answers must be a question-id to option map")
    total = len(rows)
    answered = 0
    correct_count = 0
    wrong_count = 0
    negative_per_wrong = 0.25
    checked = []
    for r in rows:
        answer = str(answers.get(str(r["id"]), answers.get(r["id"], ""))).upper().strip()
        if answer:
            answered += 1
        is_correct = answer == str(r["correct_option"]).upper()
        if is_correct:
            correct_count += 1
        elif answer:
            wrong_count += 1
        checked.append({
            "question_id": r["id"],
            "question_no": r["question_no"],
            "selected": answer,
            "correct": is_correct,
        })
    if answered < total:
        raise HTTPException(status_code=400, detail=f"Please answer all {total} MCQ questions before uploading the resume")
    raw_score = round(correct_count - (wrong_count * negative_per_wrong), 2)
    pct = round((raw_score * 100) / max(1, total), 1)
    return {
        "score": raw_score,
        "max_score": total,
        "total": total,
        "answered_count": answered,
        "correct_count": correct_count,
        "wrong_count": wrong_count,
        "negative_per_wrong": negative_per_wrong,
        "percentage": pct,
        "passed": raw_score >= 6,
        "answers": checked,
        "scoring_rule": "+1 for each correct answer, -0.25 for each wrong answer. All questions are mandatory.",
    }


@app.get("/api/demand")
def list_demand(q: str = "", status: str = "", skill: str = "", limit: int = 100, user: Dict[str, Any] = Depends(get_current_user)) -> List[Dict[str, Any]]:
    clauses = []
    values: List[Any] = []
    if q:
        clauses.append("(demand_code LIKE ? OR client_name LIKE ? OR project_name LIKE ? OR role_title LIKE ? OR domain LIKE ?)")
        like = f"%{q}%"
        values.extend([like, like, like, like, like])
    if status:
        clauses.append("status = ?")
        values.append(status)
    if skill:
        clauses.append("(required_skills LIKE ? OR role_definition LIKE ? OR role_title LIKE ?)")
        like = f"%{skill}%"
        values.extend([like, like, like])
    where = " WHERE " + " AND ".join(clauses) if clauses else ""
    values.append(limit)
    with get_db() as conn:
        rows = conn.execute(f"SELECT * FROM demand_requests {where} ORDER BY CASE priority WHEN 'Hot' THEN 0 WHEN 'High' THEN 1 WHEN 'Medium' THEN 2 ELSE 3 END, updated_at DESC LIMIT ?", values).fetchall()
    return [dict_row(r) for r in rows]


@app.get("/api/demand/{demand_id}")
def get_demand(demand_id: int, user: Dict[str, Any] = Depends(get_current_user)) -> Dict[str, Any]:
    with get_db() as conn:
        d = conn.execute("SELECT * FROM demand_requests WHERE id=?", (demand_id,)).fetchone()
        if not d:
            raise HTTPException(status_code=404, detail="Demand request not found")
        rows = conn.execute(
            """
            SELECT ds.*, c.candidate_code, c.full_name, c.email, c.phone, c.primary_skill, c.secondary_skills,
                   c.current_status, c.availability_date, c.status AS candidate_status, c.negotiated_rate, c.expected_rate,
                   c.ml_rating_score, c.fake_risk_score, c.fake_risk_level
            FROM demand_shortlists ds JOIN candidates c ON ds.candidate_id = c.id
            WHERE ds.demand_id=? ORDER BY ds.match_score DESC, ds.updated_at DESC
            """,
            (demand_id,),
        ).fetchall()
    out = dict_row(d)
    out["shortlists"] = [dict_row(r) for r in rows]
    questions = get_or_create_demand_mcqs(conn, demand_id)
    conn.commit()
    out["mcq_questions"] = safe_mcq_rows(questions, include_answers=True)
    return out




@app.get("/api/demand/{demand_id}/mcq")
def get_demand_mcq(demand_id: int, user: Dict[str, Any] = Depends(get_current_user)) -> Dict[str, Any]:
    with get_db() as conn:
        rows = get_or_create_demand_mcqs(conn, demand_id)
    return {"demand_id": demand_id, "questions": safe_mcq_rows(rows, include_answers=True)}


@app.post("/api/demand/{demand_id}/mcq/generate")
def generate_demand_mcq(demand_id: int, user: Dict[str, Any] = Depends(require_roles("Admin", "Recruiter"))) -> Dict[str, Any]:
    with get_db() as conn:
        rows = get_or_create_demand_mcqs(conn, demand_id, refresh=True)
        conn.commit()
    log_action(user["username"], "generate_demand_mcq", "demand", demand_id, "10 MCQ questions generated")
    return {"demand_id": demand_id, "questions": safe_mcq_rows(rows, include_answers=True), "message": "10 demand MCQ questions generated"}

@app.post("/api/demand")
def create_demand(payload: DemandCreate, user: Dict[str, Any] = Depends(require_roles("Admin", "Recruiter"))) -> Dict[str, Any]:
    with get_db() as conn:
        code = generate_demand_code(conn)
        cur = conn.execute(
            """
            INSERT INTO demand_requests(demand_code, client_name, project_name, role_title, role_definition, required_skills,
                domain, location, work_mode, priority, status, number_of_positions, target_customer_rate,
                max_internal_cost, start_date, duration_weeks, created_by, created_date, created_at, updated_at)
            VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
            """,
            (code, payload.client_name, payload.project_name, payload.role_title, payload.role_definition, payload.required_skills,
             payload.domain, payload.location, payload.work_mode, payload.priority, payload.status, payload.number_of_positions,
             payload.target_customer_rate, payload.max_internal_cost, payload.start_date, payload.duration_weeks, user["username"], today_date(), now_iso(), now_iso()),
        )
        conn.commit()
        demand_id = int(cur.lastrowid)
        row = conn.execute("SELECT * FROM demand_requests WHERE id=?", (demand_id,)).fetchone()
    log_action(user["username"], "create_demand", "demand", demand_id, payload.role_title)
    return dict_row(row)


@app.put("/api/demand/{demand_id}")
def update_demand(demand_id: int, payload: DemandUpdate, user: Dict[str, Any] = Depends(require_roles("Admin", "Recruiter"))) -> Dict[str, Any]:
    with get_db() as conn:
        exists = conn.execute("SELECT id FROM demand_requests WHERE id=?", (demand_id,)).fetchone()
        if not exists:
            raise HTTPException(status_code=404, detail="Demand request not found")
        conn.execute(
            """
            UPDATE demand_requests SET client_name=?, project_name=?, role_title=?, role_definition=?, required_skills=?,
                domain=?, location=?, work_mode=?, priority=?, status=?, number_of_positions=?, target_customer_rate=?,
                max_internal_cost=?, start_date=?, duration_weeks=?, updated_at=? WHERE id=?
            """,
            (payload.client_name, payload.project_name, payload.role_title, payload.role_definition, payload.required_skills,
             payload.domain, payload.location, payload.work_mode, payload.priority, payload.status, payload.number_of_positions,
             payload.target_customer_rate, payload.max_internal_cost, payload.start_date, payload.duration_weeks, now_iso(), demand_id),
        )
        conn.commit()
        row = conn.execute("SELECT * FROM demand_requests WHERE id=?", (demand_id,)).fetchone()
    log_action(user["username"], "update_demand", "demand", demand_id, payload.role_title)
    return dict_row(row)


@app.delete("/api/demand/{demand_id}")
def delete_demand(demand_id: int, user: Dict[str, Any] = Depends(require_roles("Admin", "Recruiter"))) -> Dict[str, str]:
    with get_db() as conn:
        conn.execute("DELETE FROM demand_requests WHERE id=?", (demand_id,))
        conn.commit()
    log_action(user["username"], "delete_demand", "demand", demand_id)
    return {"message": "Demand request deleted"}


@app.get("/api/demand/{demand_id}/matches")
def demand_matches(demand_id: int, limit: int = 25, user: Dict[str, Any] = Depends(get_current_user)) -> List[Dict[str, Any]]:
    with get_db() as conn:
        d = conn.execute("SELECT * FROM demand_requests WHERE id=?", (demand_id,)).fetchone()
        if not d:
            raise HTTPException(status_code=404, detail="Demand request not found")
        candidates = conn.execute("SELECT * FROM candidates ORDER BY updated_at DESC LIMIT 500").fetchall()
    demand = dict_row(d)
    matches = []
    for c in candidates:
        cd = dict_row(c)
        score = match_candidate_to_demand(cd, demand)
        item = {**cd, **score}
        matches.append(item)
    matches.sort(key=lambda x: x["match_score"], reverse=True)
    return matches[:max(1, min(limit, 100))]


@app.post("/api/demand/{demand_id}/shortlist/{candidate_id}")
def add_shortlist(demand_id: int, candidate_id: int, user: Dict[str, Any] = Depends(require_roles("Admin", "Recruiter"))) -> Dict[str, Any]:
    with get_db() as conn:
        d = conn.execute("SELECT * FROM demand_requests WHERE id=?", (demand_id,)).fetchone()
        c = conn.execute("SELECT * FROM candidates WHERE id=?", (candidate_id,)).fetchone()
        if not d:
            raise HTTPException(status_code=404, detail="Demand request not found")
        if not c:
            raise HTTPException(status_code=404, detail="Candidate not found")
        score = match_candidate_to_demand(dict_row(c), dict_row(d))
        conn.execute(
            """
            INSERT INTO demand_shortlists(demand_id, candidate_id, match_score, match_level, skill_matches, skill_gaps,
                commercial_fit, availability_fit, status, notes, created_at, updated_at)
            VALUES(?,?,?,?,?,?,?,?,?,?,?,?)
            ON CONFLICT(demand_id, candidate_id) DO UPDATE SET match_score=excluded.match_score,
                match_level=excluded.match_level, skill_matches=excluded.skill_matches, skill_gaps=excluded.skill_gaps,
                commercial_fit=excluded.commercial_fit, availability_fit=excluded.availability_fit, updated_at=excluded.updated_at
            """,
            (demand_id, candidate_id, score["match_score"], score["match_level"], ", ".join(score["skill_matches"]),
             ", ".join(score["skill_gaps"]), score["commercial_fit"], score.get("availability_fit", ""), "Shortlisted", "", now_iso(), now_iso()),
        )
        conn.commit()
        row = conn.execute("SELECT * FROM demand_shortlists WHERE demand_id=? AND candidate_id=?", (demand_id, candidate_id)).fetchone()
    log_action(user["username"], "shortlist_candidate", "demand", demand_id, f"candidate={candidate_id}; score={score['match_score']}")
    return dict_row(row)


@app.patch("/api/demand/shortlist/{shortlist_id}")
def update_shortlist(shortlist_id: int, payload: ShortlistUpdate, user: Dict[str, Any] = Depends(require_roles("Admin", "Recruiter"))) -> Dict[str, Any]:
    with get_db() as conn:
        row = conn.execute("SELECT * FROM demand_shortlists WHERE id=?", (shortlist_id,)).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Shortlist item not found")
        conn.execute("UPDATE demand_shortlists SET status=?, notes=?, updated_at=? WHERE id=?", (payload.status, payload.notes, now_iso(), shortlist_id))
        conn.commit()
        updated = conn.execute("SELECT * FROM demand_shortlists WHERE id=?", (shortlist_id,)).fetchone()
    log_action(user["username"], "update_shortlist", "shortlist", shortlist_id, payload.status)
    return dict_row(updated)


@app.post("/api/resume-links")
def create_resume_link(payload: LinkCreate, user: Dict[str, Any] = Depends(require_roles("Admin", "Recruiter"))) -> Dict[str, Any]:
    token = secrets.token_urlsafe(32)
    expires_at = (datetime.utcnow() + timedelta(hours=48)).replace(microsecond=0).isoformat() + "Z"
    with get_db() as conn:
        if payload.candidate_id:
            exists = conn.execute("SELECT id FROM candidates WHERE id=?", (payload.candidate_id,)).fetchone()
            if not exists:
                raise HTTPException(status_code=404, detail="Candidate not found")
        demand = None
        if payload.demand_id:
            demand = conn.execute("SELECT * FROM demand_requests WHERE id=?", (payload.demand_id,)).fetchone()
            if not demand:
                raise HTTPException(status_code=404, detail="Demand record not found")
        role_title = (payload.role_title or "").strip() or (demand["role_title"] if demand else "")
        role_definition = (payload.role_definition or "").strip() or (demand["role_definition"] if demand else "")
        include_mcq_enabled = 1 if (payload.demand_id and payload.include_mcq) else 0
        cur = conn.execute(
            """
            INSERT INTO public_upload_links(token, role_title, role_definition, candidate_id, demand_id, include_mcq, created_by, expires_at, is_active, created_at)
            VALUES(?,?,?,?,?,?,?,?,?,?)
            """,
            (token, role_title, role_definition, payload.candidate_id, payload.demand_id, include_mcq_enabled, user["username"], expires_at, 1, now_iso()),
        )
        conn.commit()
        link_id = cur.lastrowid
    detail = f"expires={expires_at}; demand_id={payload.demand_id or ''}; candidate_id={payload.candidate_id or ''}; include_mcq={include_mcq_enabled}"
    log_action(user["username"], "create_public_resume_link", "public_upload_link", link_id, detail)
    return {
        "id": link_id,
        "token": token,
        "role_title": role_title,
        "role_definition": role_definition,
        "candidate_id": payload.candidate_id,
        "demand_id": payload.demand_id,
        "include_mcq": bool(include_mcq_enabled),
        "expires_at": expires_at,
        "upload_path": f"/public-upload/{token}",
    }


@app.get("/api/resume-links")
def list_resume_links(user: Dict[str, Any] = Depends(require_roles("Admin", "Recruiter"))) -> List[Dict[str, Any]]:
    with get_db() as conn:
        rows = conn.execute(
            """
            SELECT pul.*, dr.demand_code, dr.client_name, dr.project_name, dr.role_title AS demand_role_title
            FROM public_upload_links pul
            LEFT JOIN demand_requests dr ON dr.id = pul.demand_id
            ORDER BY pul.created_at DESC LIMIT 100
            """
        ).fetchall()
    result = []
    for r in rows:
        item = dict_row(r)
        item["is_expired"] = parse_iso(item["expires_at"]) < datetime.utcnow()
        item["upload_path"] = f"/public-upload/{item['token']}"
        item["demand_label"] = ""
        if item.get("demand_id"):
            parts = [item.get("demand_code") or f"Demand #{item['demand_id']}", item.get("client_name") or "", item.get("demand_role_title") or ""]
            item["demand_label"] = " · ".join([p for p in parts if p])
        result.append(item)
    return result


@app.post("/api/resume-links/{link_id}/revoke")
def revoke_resume_link(link_id: int, user: Dict[str, Any] = Depends(require_roles("Admin", "Recruiter"))) -> Dict[str, str]:
    with get_db() as conn:
        row = conn.execute("SELECT id FROM public_upload_links WHERE id=?", (link_id,)).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Public upload link not found")
        conn.execute("UPDATE public_upload_links SET is_active=0, revoked_at=?, revoked_by=? WHERE id=?", (now_iso(), user["username"], link_id))
        conn.commit()
    log_action(user["username"], "revoke_public_resume_link", "public_upload_link", link_id)
    return {"message": "Public upload link revoked"}


@app.get("/api/public-upload/{token}")
def public_link_info(token: str) -> Dict[str, Any]:
    assert_public_token(token)
    with get_db() as conn:
        row = conn.execute("SELECT * FROM public_upload_links WHERE token=?", (token,)).fetchone()
        candidate = None
        demand = None
        if row and row["candidate_id"]:
            candidate = conn.execute("SELECT id, candidate_code, full_name, email FROM candidates WHERE id=?", (row["candidate_id"],)).fetchone()
        questions = []
        if row and row["demand_id"]:
            demand = conn.execute("SELECT id, demand_code, client_name, project_name, role_title, role_definition, required_skills, start_date FROM demand_requests WHERE id=?", (row["demand_id"],)).fetchone()
            if int(row["include_mcq"] or 0) == 1:
                questions = get_or_create_demand_mcqs(conn, int(row["demand_id"]))
                conn.commit()
    if not row:
        raise HTTPException(status_code=404, detail="Upload link not found")
    item = dict_row(row)
    if item["is_active"] != 1 or item.get("revoked_at") or parse_iso(item["expires_at"]) < datetime.utcnow():
        raise HTTPException(status_code=410, detail="Upload link has expired or has been revoked")
    demand_dict = dict_row(demand) if demand else None
    is_update_only = bool(item.get("used_at") and item.get("candidate_id"))
    mcq_enabled = bool(int(item.get("include_mcq") or 0) == 1 and item.get("demand_id") and not is_update_only)
    return {
        "role_title": item["role_title"] or (demand_dict.get("role_title") if demand_dict else ""),
        "role_definition": item["role_definition"] or (demand_dict.get("role_definition") if demand_dict else ""),
        "demand_id": item.get("demand_id"),
        "demand": demand_dict,
        "expires_at": item["expires_at"],
        "used_at": item.get("used_at"),
        "use_count": int(item.get("use_count") or 0),
        "mode": "update_resume_only" if is_update_only else "first_upload",
        "candidate": dict_row(candidate) if candidate else None,
        "include_mcq": bool(int(item.get("include_mcq") or 0) == 1),
        "mcq_required": mcq_enabled,
        "mcq_questions": safe_mcq_rows(questions, include_answers=False) if mcq_enabled else [],
        "message": "This link has already created the candidate record. You can only upload an updated resume for the same candidate." if is_update_only else ("This link will create or bind the candidate record on first upload. MCQ screening is enabled for this demand-linked link." if mcq_enabled else "This link will create or bind the candidate record on first upload. MCQ screening is not required for this link."),
    }



@app.post("/api/public-upload/{token}/mcq-score")
def public_score_mcq(token: str, payload: PublicMcqSubmit) -> Dict[str, Any]:
    assert_public_token(token)
    with get_db() as conn:
        link = conn.execute("SELECT * FROM public_upload_links WHERE token=?", (token,)).fetchone()
        if not link:
            raise HTTPException(status_code=404, detail="Upload link not found")
        item = dict_row(link)
        if item["is_active"] != 1 or item.get("revoked_at") or parse_iso(item["expires_at"]) < datetime.utcnow():
            raise HTTPException(status_code=410, detail="Upload link has expired or has been revoked")
        if item.get("used_at") and item.get("candidate_id"):
            raise HTTPException(status_code=400, detail="This link is already used. MCQ is not required for resume-only updates.")
        if not item.get("demand_id"):
            raise HTTPException(status_code=400, detail="This public upload link is not connected to a demand record")
        if int(item.get("include_mcq") or 0) != 1:
            raise HTTPException(status_code=400, detail="MCQ is disabled for this public upload link")
        result = score_mcq_answers(conn, int(item["demand_id"]), json.dumps({"answers": payload.answers}))
    return {
        "score": result["score"],
        "max_score": result["max_score"],
        "total": result["total"],
        "answered_count": result["answered_count"],
        "correct_count": result["correct_count"],
        "wrong_count": result["wrong_count"],
        "negative_per_wrong": result["negative_per_wrong"],
        "percentage": result["percentage"],
        "passed": result["passed"],
        "scoring_rule": result["scoring_rule"],
        "message": "MCQ completed. You may now upload your resume.",
    }


@app.post("/api/public-upload/{token}")
def public_upload_resume(token: str, role_title: str = Form(""), role_definition: str = Form(""), full_name: str = Form(""), email: str = Form(""), phone: str = Form(""), location: str = Form(""), current_company: str = Form(""), previous_companies: str = Form(""), project_details: str = Form(""), available_by_date: str = Form(""), notice_period_days: int = Form(0), mcq_answers: str = Form(""), photograph: Optional[UploadFile] = File(None), file: UploadFile = File(...)) -> Dict[str, Any]:
    assert_public_token(token)
    with get_db() as conn:
        link = conn.execute("SELECT * FROM public_upload_links WHERE token=?", (token,)).fetchone()
        if not link:
            raise HTTPException(status_code=404, detail="Upload link not found")
        link_dict = dict_row(link)
        if link_dict["is_active"] != 1 or link_dict.get("revoked_at") or parse_iso(link_dict["expires_at"]) < datetime.utcnow():
            raise HTTPException(status_code=410, detail="Upload link has expired or has been revoked")
        demand = conn.execute("SELECT * FROM demand_requests WHERE id=?", (link_dict.get("demand_id"),)).fetchone() if link_dict.get("demand_id") else None
        demand_dict = dict_row(demand) if demand else None
    link_is_used = bool(link_dict.get("used_at") and link_dict.get("candidate_id"))
    final_role_title = (
        link_dict.get("role_title") if link_is_used else (role_title or link_dict.get("role_title") or (demand_dict.get("role_title") if demand_dict else ""))
    ) or "Candidate Uploaded Resume"
    final_role_def = (
        link_dict.get("role_definition") if link_is_used else (role_definition or link_dict.get("role_definition") or (demand_dict.get("role_definition") if demand_dict else ""))
    ) or ""
    mcq_result = None
    mcq_enabled = bool(demand_dict and not link_is_used and int(link_dict.get("include_mcq") or 0) == 1)
    if mcq_enabled:
        with get_db() as conn:
            mcq_result = score_mcq_answers(conn, int(demand_dict["id"]), mcq_answers)
    if not link_is_used and (photograph is None or not getattr(photograph, "filename", "")):
        raise HTTPException(status_code=400, detail="Please upload a recent candidate photograph")
    photo_name = ""
    photo_path = ""
    if photograph is not None and getattr(photograph, "filename", ""):
        photo_name, photo_path = save_photo_upload(photograph, f"photo_{token[:8]}")
    file_name, file_path = save_upload(file, f"public_{token[:8]}")
    resume_text = extract_text_from_file(file_path, file_name)
    analysis = analyze_resume_text(resume_text, final_role_def, final_role_title)
    shortlist_info = None
    with get_db() as conn:
        candidate_id = link_dict.get("candidate_id")
        first_use = not bool(candidate_id)
        if not candidate_id:
            code = generate_candidate_code(conn)
            derived_name = full_name or analysis.get("detected_name") or "Candidate"
            cur = conn.execute(
                """
                INSERT INTO candidates(candidate_code, full_name, email, phone, location, current_status, availability_date, available_by_date, notice_period_days,
                    employment_type, source, recruiter_owner, total_experience, relevant_experience, primary_skill,
                    secondary_skills, domain_exposure, proficiency, certifications, portfolio_url, current_company, previous_companies, project_details, photo_file_name, photo_file_path, expected_rate, negotiated_rate,
                    internal_level, resume_text, status, created_date, created_at, updated_at)
                VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
                """,
                (code, derived_name, email or analysis.get("email", ""), phone or analysis.get("phone", ""), location,
                 "Screening", "To be confirmed", available_by_date, int(notice_period_days or 0), "Contract", "Public Upload Link", link_dict.get("created_by", "Recruiter"),
                 analysis.get("experience_years", 0), analysis.get("experience_years", 0), "", "", "", "Intermediate", "", "",
                 current_company, previous_companies, project_details, photo_name, photo_path, 0, 0,
                 "L2 - Mid-level", "", "Screening", today_date(), now_iso(), now_iso()),
            )
            candidate_id = int(cur.lastrowid)
        else:
            # Once used, the public URL is locked to the first candidate. Ignore identity fields on later uploads.
            exists = conn.execute("SELECT id FROM candidates WHERE id=?", (candidate_id,)).fetchone()
            if not exists:
                raise HTTPException(status_code=409, detail="The candidate linked to this public URL no longer exists")
        if photo_name or (not first_use and (current_company or previous_companies or project_details)):
            conn.execute(
                """
                UPDATE candidates SET
                    current_company=COALESCE(NULLIF(?,''), current_company),
                    previous_companies=COALESCE(NULLIF(?,''), previous_companies),
                    project_details=COALESCE(NULLIF(?,''), project_details),
                    photo_file_name=COALESCE(NULLIF(?,''), photo_file_name),
                    photo_file_path=COALESCE(NULLIF(?,''), photo_file_path),
                    updated_at=?
                WHERE id=?
                """,
                (current_company, previous_companies, project_details, photo_name, photo_path, now_iso(), int(candidate_id)),
            )
        resume_id = create_resume_record(conn, int(candidate_id), file_name, file_path, resume_text, final_role_title, final_role_def, "Public Candidate Upload", analysis, update_identity=first_use)
        if mcq_result and link_dict.get("demand_id"):
            conn.execute(
                """
                INSERT INTO public_mcq_results(public_link_id, demand_id, candidate_id, score, total, answered_count, correct_count, wrong_count, negative_per_wrong, percentage, passed, answers_json, created_at)
                VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)
                """,
                (int(link_dict["id"]), int(link_dict["demand_id"]), int(candidate_id), float(mcq_result["score"]), int(mcq_result["total"]),
                 int(mcq_result["answered_count"]), int(mcq_result["correct_count"]), int(mcq_result["wrong_count"]),
                 float(mcq_result["negative_per_wrong"]), float(mcq_result["percentage"]),
                 1 if mcq_result.get("passed") else 0, json.dumps(mcq_result), now_iso()),
            )
        if link_dict.get("demand_id"):
            demand_row = conn.execute("SELECT * FROM demand_requests WHERE id=?", (link_dict.get("demand_id"),)).fetchone()
            candidate_row = conn.execute("SELECT * FROM candidates WHERE id=?", (candidate_id,)).fetchone()
            if demand_row and candidate_row:
                score = match_candidate_to_demand(dict_row(candidate_row), dict_row(demand_row))
                conn.execute(
                    """
                    INSERT INTO demand_shortlists(demand_id, candidate_id, match_score, match_level, skill_matches, skill_gaps,
                        commercial_fit, availability_fit, status, notes, created_at, updated_at)
                    VALUES(?,?,?,?,?,?,?,?,?,?,?,?)
                    ON CONFLICT(demand_id, candidate_id) DO UPDATE SET match_score=excluded.match_score,
                        match_level=excluded.match_level, skill_matches=excluded.skill_matches, skill_gaps=excluded.skill_gaps,
                        commercial_fit=excluded.commercial_fit, availability_fit=excluded.availability_fit,
                        status=excluded.status, notes=excluded.notes, updated_at=excluded.updated_at
                    """,
                    (int(link_dict["demand_id"]), int(candidate_id), score["match_score"], score["match_level"], ", ".join(score["skill_matches"]),
                     ", ".join(score["skill_gaps"]), score["commercial_fit"], score.get("availability_fit", ""),
                     "Resume Uploaded", f"Auto-linked from public upload link; MCQ {mcq_result['score']}/{mcq_result['max_score']} ({mcq_result['correct_count']} correct, {mcq_result['wrong_count']} wrong, -{mcq_result['negative_per_wrong']} per wrong)" if mcq_result else "Auto-linked from public upload link", now_iso(), now_iso()),
                )
                shortlist_info = {"demand_id": int(link_dict["demand_id"]), **score}
        if first_use:
            conn.execute(
                "UPDATE public_upload_links SET candidate_id=?, used_at=?, last_uploaded_at=?, use_count=COALESCE(use_count,0)+1, locked_note=? WHERE token=?",
                (candidate_id, now_iso(), now_iso(), "Locked after first upload; future use updates resume only", token),
            )
        else:
            conn.execute(
                "UPDATE public_upload_links SET last_uploaded_at=?, use_count=COALESCE(use_count,0)+1 WHERE token=?",
                (now_iso(), token),
            )
        conn.commit()
    detail = f"resume_id={resume_id}; fit={analysis['fit_score']}; risk={analysis['fake_risk_score']}; demand_id={link_dict.get('demand_id') or ''}"
    log_action("Public Candidate", "public_resume_upload_update_only" if not first_use else "public_resume_first_upload", "candidate", int(candidate_id), detail)
    return {"message": "Resume uploaded successfully", "mode": "created_candidate" if first_use else "updated_resume_only", "candidate_id": candidate_id, "resume_id": resume_id, "analysis": analysis, "shortlist": shortlist_info, "mcq_result": mcq_result}



def month_key(value: Optional[str]) -> str:
    if not value:
        return "Unknown"
    try:
        return parse_iso(value).strftime("%Y-%m")
    except Exception:
        return str(value)[:7] if len(str(value)) >= 7 else "Unknown"


def record_created_period(row: Dict[str, Any]) -> str:
    # Prefer explicit created_date for business trend analytics; fall back to system timestamp for older records.
    return month_key(row.get("created_date") or row.get("created_at"))


def normalize_skill_tokens(text_value: str) -> List[str]:
    if not text_value:
        return []
    found = extract_skills(text_value)
    if found:
        return found
    return [t.strip().title() for t in re.split(r"[,;/|]", text_value) if t.strip()][:8]


def top_counter_rows(counter: Counter, limit: int = 10, name_key: str = "skill", value_key: str = "count") -> List[Dict[str, Any]]:
    return [{name_key: name, value_key: count} for name, count in counter.most_common(limit)]


@app.get("/api/intelligence/candidate/{candidate_id}/role-suitability")
def candidate_role_suitability(candidate_id: int, user: Dict[str, Any] = Depends(get_current_user)) -> Dict[str, Any]:
    with get_db() as conn:
        c = conn.execute("SELECT * FROM candidates WHERE id=?", (candidate_id,)).fetchone()
        if not c:
            raise HTTPException(status_code=404, detail="Candidate not found")
        demands = conn.execute("SELECT * FROM demand_requests WHERE status NOT IN ('Closed','Cancelled') ORDER BY updated_at DESC").fetchall()
    candidate = dict_row(c)
    matches = []
    for d in demands:
        demand = dict_row(d)
        score = match_candidate_to_demand(candidate, demand)
        matches.append({
            "demand_id": demand["id"],
            "demand_code": demand["demand_code"],
            "client_name": demand.get("client_name") or "",
            "project_name": demand.get("project_name") or "",
            "role_title": demand.get("role_title") or "",
            "required_skills": demand.get("required_skills") or "",
            "priority": demand.get("priority") or "",
            "status": demand.get("status") or "",
            **score,
        })
    matches.sort(key=lambda x: x.get("match_score", 0), reverse=True)
    return {"candidate": candidate, "matches": matches[:25]}


@app.get("/api/intelligence/demand/{demand_id}/candidate-shortlist")
def demand_candidate_shortlist(demand_id: int, user: Dict[str, Any] = Depends(get_current_user)) -> Dict[str, Any]:
    with get_db() as conn:
        d = conn.execute("SELECT * FROM demand_requests WHERE id=?", (demand_id,)).fetchone()
        if not d:
            raise HTTPException(status_code=404, detail="Demand request not found")
        candidates = conn.execute("SELECT * FROM candidates ORDER BY updated_at DESC").fetchall()
    demand = dict_row(d)
    matches = []
    for c in candidates:
        candidate = dict_row(c)
        score = match_candidate_to_demand(candidate, demand)
        matches.append({
            "candidate_id": candidate["id"],
            "candidate_code": candidate.get("candidate_code"),
            "full_name": candidate.get("full_name"),
            "email": candidate.get("email"),
            "primary_skill": candidate.get("primary_skill"),
            "secondary_skills": candidate.get("secondary_skills"),
            "current_status": candidate.get("current_status"),
            "total_experience": candidate.get("total_experience"),
            "relevant_experience": candidate.get("relevant_experience"),
            "available_by_date": candidate.get("available_by_date"),
            "notice_period_days": candidate.get("notice_period_days"),
            "expected_rate": candidate.get("expected_rate"),
            "negotiated_rate": candidate.get("negotiated_rate"),
            "ml_rating_score": candidate.get("ml_rating_score"),
            "fake_risk_score": candidate.get("fake_risk_score"),
            **score,
        })
    matches.sort(key=lambda x: x.get("match_score", 0), reverse=True)
    return {"demand": demand, "matches": matches[:50]}


@app.get("/api/intelligence/trends")
def intelligence_trends(user: Dict[str, Any] = Depends(get_current_user)) -> Dict[str, Any]:
    with get_db() as conn:
        demands = [dict_row(r) for r in conn.execute("SELECT * FROM demand_requests").fetchall()]
        candidates = [dict_row(r) for r in conn.execute("SELECT * FROM candidates").fetchall()]

    demand_by_month: Counter = Counter()
    supply_by_month: Counter = Counter()
    demand_skill_counter: Counter = Counter()
    supply_skill_counter: Counter = Counter()

    for d in demands:
        demand_by_month[record_created_period(d)] += int(d.get("number_of_positions") or 1)
        for skill in normalize_skill_tokens(" ".join([str(d.get("role_title") or ""), str(d.get("required_skills") or ""), str(d.get("role_definition") or "")])):
            demand_skill_counter[skill] += int(d.get("number_of_positions") or 1)

    for c in candidates:
        supply_by_month[record_created_period(c)] += 1
        for skill in normalize_skill_tokens(" ".join([str(c.get("primary_skill") or ""), str(c.get("secondary_skills") or ""), str(c.get("resume_text") or "")])):
            supply_skill_counter[skill] += 1

    months = sorted(set(demand_by_month.keys()) | set(supply_by_month.keys()))
    if "Unknown" in months:
        months = [m for m in months if m != "Unknown"] + ["Unknown"]

    skill_gap_rows = []
    all_skills = set(demand_skill_counter.keys()) | set(supply_skill_counter.keys())
    for skill in all_skills:
        demand_count = int(demand_skill_counter.get(skill, 0))
        supply_count = int(supply_skill_counter.get(skill, 0))
        gap = demand_count - supply_count
        skill_gap_rows.append({"skill": skill, "demand": demand_count, "supply": supply_count, "gap": gap})
    skill_gap_rows.sort(key=lambda r: (r["gap"], r["demand"]), reverse=True)

    return {
        "demand_trend": [{"period": m, "demand": int(demand_by_month.get(m, 0))} for m in months],
        "supply_trend": [{"period": m, "supply": int(supply_by_month.get(m, 0))} for m in months],
        "combined_trend": [{"period": m, "demand": int(demand_by_month.get(m, 0)), "supply": int(supply_by_month.get(m, 0))} for m in months],
        "top_demand_skills": top_counter_rows(demand_skill_counter, 12, "skill", "demand"),
        "top_supply_skills": top_counter_rows(supply_skill_counter, 12, "skill", "supply"),
        "skill_gaps": skill_gap_rows[:15],
    }


@app.get("/api/intelligence/market-signals")
def market_signals(user: Dict[str, Any] = Depends(get_current_user)) -> Dict[str, Any]:
    trends = intelligence_trends(user)
    gaps = trends.get("skill_gaps", [])[:8]
    top_demand = trends.get("top_demand_skills", [])[:8]
    research = []
    for item in gaps:
        skill = item.get("skill")
        if not skill:
            continue
        q1 = quote_plus(f"{skill} hiring demand trend India 2026")
        q2 = quote_plus(f"{skill} talent shortage India 2026")
        q3 = quote_plus(f"{skill} salary trend India developer 2026")
        research.append({
            "skill": skill,
            "internal_demand": item.get("demand", 0),
            "internal_supply": item.get("supply", 0),
            "internal_gap": item.get("gap", 0),
            "google_trend_search": f"https://www.google.com/search?q={q1}",
            "talent_shortage_search": f"https://www.google.com/search?q={q2}",
            "salary_trend_search": f"https://www.google.com/search?q={q3}",
        })

    recommendations = []
    for item in gaps[:5]:
        if item.get("gap", 0) > 0:
            recommendations.append(f"Build or source more {item['skill']} profiles: internal demand is {item['demand']} against supply of {item['supply']}.")
    if not recommendations and top_demand:
        recommendations.append("Demand and supply look broadly balanced in the current internal database. Continue monitoring external market trends before committing bench hiring.")
    recommendations.append("Use the internet research links as validation signals before making bench hiring decisions. The app does not scrape job boards automatically in this local build.")

    return {
        "summary": "Market signal view combines internal demand/supply gaps with external internet research links for recruiter validation.",
        "research_links": research,
        "recommendations": recommendations,
        "note": "Live web scraping/API integration can be added later using a compliant jobs-data provider or search API.",
    }

@app.get("/api/ml/skills-analytics")
def skills_analytics(user: Dict[str, Any] = Depends(get_current_user)) -> Dict[str, Any]:
    with get_db() as conn:
        candidates = conn.execute("SELECT primary_skill, secondary_skills, ml_rating_score, fake_risk_score FROM candidates").fetchall()
        resumes = conn.execute("SELECT role_title, role_definition, skill_matches, skill_gaps, fit_score, fake_risk_score FROM resumes").fetchall()
        demand_requests = conn.execute("SELECT role_title, role_definition, required_skills, status FROM demand_requests").fetchall()
    supply = Counter()
    for c in candidates:
        text = f"{c['primary_skill'] or ''}, {c['secondary_skills'] or ''}"
        for s in extract_skills(text):
            supply[s] += 1
    demand = Counter()
    for d in demand_requests:
        if d["status"] not in ["Closed", "Cancelled", "On Hold"]:
            for s in extract_skills(f"{d['role_title'] or ''} {d['required_skills'] or ''} {d['role_definition'] or ''}"):
                demand[s] += 1
    # Resume role definitions are also used as weak demand signals because they show recent recruiter requests.
    for r in resumes:
        for s in extract_skills(f"{r['role_title'] or ''} {r['role_definition'] or ''}"):
            demand[s] += 0.35
    clusters = []
    for cluster, skills in SKILL_CLUSTERS.items():
        cluster_supply = sum(supply.get(s, 0) for s in skills)
        cluster_demand = sum(demand.get(s, 0) for s in skills)
        clusters.append({"cluster": cluster, "supply": cluster_supply, "demand": cluster_demand, "gap": max(0, cluster_demand - cluster_supply)})
    scarce = []
    for skill in sorted(set(list(supply.keys()) + list(demand.keys()))):
        scarce.append({"skill": skill, "supply": supply.get(skill, 0), "demand": demand.get(skill, 0), "gap": max(0, demand.get(skill, 0) - supply.get(skill, 0))})
    scarce = sorted(scarce, key=lambda x: (x["gap"], x["demand"]), reverse=True)[:12]
    total_profiles = len(candidates)
    avg_fit = round(sum(int(r["fit_score"] or 0) for r in resumes) / max(1, len(resumes)), 1)
    high_risk = sum(1 for r in resumes if int(r["fake_risk_score"] or 0) >= 75)
    recommendations = []
    if any(x["gap"] > 0 for x in scarce):
        recommendations.append("Prioritize sourcing for skill gaps where role demand is higher than available screened supply.")
    if high_risk:
        recommendations.append(f"Run manual validation for {high_risk} resume(s) with high fake-resume risk.")
    if avg_fit < 65 and resumes:
        recommendations.append("Average fit score is below 65; tighten role definitions or improve candidate sourcing quality.")
    if not recommendations:
        recommendations.append("Skill supply and resume quality are currently healthy based on available data.")
    return {
        "total_profiles": total_profiles,
        "total_resumes": len(resumes),
        "average_resume_fit": avg_fit,
        "high_risk_resumes": high_risk,
        "top_supply_skills": [{"skill": k, "count": v} for k, v in supply.most_common(12)],
        "open_demand_requests": sum(1 for d in demand_requests if d["status"] not in ["Closed", "Cancelled", "On Hold"]),
        "top_demand_skills": [{"skill": k, "count": round(v, 2)} for k, v in demand.most_common(12)],
        "clusters": clusters,
        "skill_gaps": scarce,
        "recommendations": recommendations,
    }


@app.get("/api/export/candidates.csv")
def export_candidates(user: Dict[str, Any] = Depends(get_current_user)):
    with get_db() as conn:
        rows = conn.execute("SELECT * FROM candidates ORDER BY updated_at DESC").fetchall()
    output = io.StringIO()
    if rows:
        fieldnames = rows[0].keys()
    else:
        fieldnames = ["candidate_code", "full_name", "email", "primary_skill", "status", "ml_rating_score", "fake_risk_score"]
    writer = csv.DictWriter(output, fieldnames=fieldnames)
    writer.writeheader()
    for row in rows:
        writer.writerow(dict_row(row))
    output.seek(0)
    return StreamingResponse(iter([output.getvalue()]), media_type="text/csv", headers={"Content-Disposition": "attachment; filename=1resource_resume_bank.csv"})


@app.get("/api/logs")
def logs(user: Dict[str, Any] = Depends(require_roles("Admin"))) -> List[Dict[str, Any]]:
    with get_db() as conn:
        rows = conn.execute("SELECT * FROM activity_logs ORDER BY created_at DESC LIMIT 100").fetchall()
    return [dict_row(r) for r in rows]


@app.get("/api/security/status")
def security_status(user: Dict[str, Any] = Depends(require_roles("Admin"))) -> Dict[str, Any]:
    with get_db() as conn:
        active_sessions = conn.execute("SELECT COUNT(*) AS c FROM sessions").fetchone()["c"]
        locked_users = conn.execute("SELECT COUNT(*) AS c FROM users WHERE locked_until IS NOT NULL AND locked_until > ?", (now_iso(),)).fetchone()["c"]
        active_public_links = conn.execute("SELECT COUNT(*) AS c FROM public_upload_links WHERE is_active=1 AND expires_at > ?", (now_iso(),)).fetchone()["c"]
        high_risk = conn.execute("SELECT COUNT(*) AS c FROM candidates WHERE fake_risk_score >= 75").fetchone()["c"]
    return {
        "version": APP_VERSION,
        "session_ttl_hours": SESSION_TTL_HOURS,
        "allowed_origins": ALLOWED_ORIGINS,
        "max_upload_mb": MAX_UPLOAD_MB,
        "allowed_upload_extensions": sorted(ALLOWED_UPLOAD_EXTENSIONS),
        "active_sessions": active_sessions,
        "locked_users": locked_users,
        "active_public_links": active_public_links,
        "high_risk_candidates": high_risk,
        "security_controls": [
            "Restricted CORS origins",
            "HTTP security headers",
            "Session expiry",
            "Login lockout after repeated failed attempts",
            "Upload file type and size validation",
            "Public-link expiry and revoke control",
            "Public-link repeat use locked to resume-only update",
            "Demand-linked MCQ screening before first public resume upload",
            "Authenticated resume downloads",
            "Audit logs for critical actions",
        ],
    }



@app.get("/api/potential-clients")
def list_potential_clients(user: Dict[str, Any] = Depends(require_roles("Admin"))) -> List[Dict[str, Any]]:
    with get_db() as conn:
        rows = conn.execute("SELECT * FROM potential_clients ORDER BY updated_at DESC, id DESC").fetchall()
    return [dict_row(r) for r in rows]


@app.post("/api/potential-clients")
def create_potential_client(payload: PotentialClientCreate, user: Dict[str, Any] = Depends(require_roles("Admin"))) -> Dict[str, Any]:
    company_name = (payload.company_name or "").strip()
    email = validate_outreach_email(payload.email)
    if not company_name:
        raise HTTPException(status_code=400, detail="Company name is required")
    with get_db() as conn:
        cur = conn.execute(
            """
            INSERT INTO potential_clients(company_name, contact_name, email, phone, segment, status, notes, created_by, created_at, updated_at)
            VALUES(?,?,?,?,?,?,?,?,?,?)
            """,
            (company_name, (payload.contact_name or "").strip(), email, (payload.phone or "").strip(),
             (payload.segment or "").strip(), (payload.status or "Prospect").strip(), (payload.notes or "").strip(),
             user["username"], now_iso(), now_iso()),
        )
        conn.commit()
        row = conn.execute("SELECT * FROM potential_clients WHERE id=?", (cur.lastrowid,)).fetchone()
    log_action(user["username"], "create_potential_client", "potential_client", cur.lastrowid, company_name)
    return dict_row(row)


@app.put("/api/potential-clients/{client_id}")
def update_potential_client(client_id: int, payload: PotentialClientUpdate, user: Dict[str, Any] = Depends(require_roles("Admin"))) -> Dict[str, Any]:
    company_name = (payload.company_name or "").strip()
    email = validate_outreach_email(payload.email)
    if not company_name:
        raise HTTPException(status_code=400, detail="Company name is required")
    with get_db() as conn:
        existing = conn.execute("SELECT id FROM potential_clients WHERE id=?", (client_id,)).fetchone()
        if not existing:
            raise HTTPException(status_code=404, detail="Potential client not found")
        conn.execute(
            """
            UPDATE potential_clients
            SET company_name=?, contact_name=?, email=?, phone=?, segment=?, status=?, notes=?, updated_at=?
            WHERE id=?
            """,
            (company_name, (payload.contact_name or "").strip(), email, (payload.phone or "").strip(),
             (payload.segment or "").strip(), (payload.status or "Prospect").strip(), (payload.notes or "").strip(),
             now_iso(), client_id),
        )
        conn.commit()
        row = conn.execute("SELECT * FROM potential_clients WHERE id=?", (client_id,)).fetchone()
    log_action(user["username"], "update_potential_client", "potential_client", client_id, company_name)
    return dict_row(row)


@app.delete("/api/potential-clients/{client_id}")
def delete_potential_client(client_id: int, user: Dict[str, Any] = Depends(require_roles("Admin"))) -> Dict[str, str]:
    with get_db() as conn:
        conn.execute("DELETE FROM potential_clients WHERE id=?", (client_id,))
        conn.commit()
    log_action(user["username"], "delete_potential_client", "potential_client", client_id)
    return {"message": "Potential client deleted"}


@app.get("/api/outreach/users")
def list_outreach_users(user: Dict[str, Any] = Depends(require_roles("Admin"))) -> List[Dict[str, Any]]:
    with get_db() as conn:
        rows = conn.execute(
            "SELECT id, username, full_name, email, phone, role, is_active FROM users WHERE COALESCE(email,'') <> '' ORDER BY full_name, username"
        ).fetchall()
    return [dict_row(r) for r in rows]


@app.get("/api/outreach/logs")
def list_outreach_logs(user: Dict[str, Any] = Depends(require_roles("Admin"))) -> List[Dict[str, Any]]:
    with get_db() as conn:
        rows = conn.execute("SELECT * FROM email_outreach_logs ORDER BY created_at DESC, id DESC LIMIT 200").fetchall()
    return [dict_row(r) for r in rows]


@app.get("/api/outreach/smtp-status")
def outreach_smtp_status(user: Dict[str, Any] = Depends(require_roles("Admin"))) -> Dict[str, Any]:
    return email_provider_status()


@app.post("/api/outreach/test-smtp")
def outreach_test_smtp(user: Dict[str, Any] = Depends(require_roles("Admin"))) -> Dict[str, Any]:
    result = test_email_connection()
    log_action(user["username"], "test_smtp_connection", "email_outreach", None, result.get("message", ""))
    return result


@app.post("/api/outreach/send")
def send_customer_outreach(payload: OutreachEmailRequest, user: Dict[str, Any] = Depends(require_roles("Admin"))) -> Dict[str, Any]:
    email_type = validate_email_type(payload.email_type)
    subject = (payload.subject or "").strip()
    body = (payload.body or "").strip()
    if not subject:
        raise HTTPException(status_code=400, detail="Subject is required")
    if not body:
        raise HTTPException(status_code=400, detail="Email body is required")
    if not payload.client_ids and not payload.user_ids:
        raise HTTPException(status_code=400, detail="Select at least one potential client or user email")

    recipients: List[Dict[str, Any]] = []
    with get_db() as conn:
        if payload.client_ids:
            placeholders = ",".join(["?"] * len(payload.client_ids))
            client_rows = conn.execute(f"SELECT * FROM potential_clients WHERE id IN ({placeholders})", tuple(payload.client_ids)).fetchall()
            for r in client_rows:
                recipients.append({
                    "recipient_type": "client",
                    "client_id": r["id"],
                    "user_id": None,
                    "recipient_name": r["contact_name"] or r["company_name"],
                    "recipient_email": r["email"],
                })
        if payload.user_ids:
            placeholders = ",".join(["?"] * len(payload.user_ids))
            user_rows = conn.execute(f"SELECT id, full_name, username, email FROM users WHERE id IN ({placeholders}) AND COALESCE(email,'') <> ''", tuple(payload.user_ids)).fetchall()
            for r in user_rows:
                recipients.append({
                    "recipient_type": "user",
                    "client_id": None,
                    "user_id": r["id"],
                    "recipient_name": r["full_name"] or r["username"],
                    "recipient_email": r["email"],
                })

        if not recipients:
            raise HTTPException(status_code=400, detail="No valid recipient emails found")

        sent = queued = failed = 0
        results = []
        for rec in recipients:
            recipient_email = validate_outreach_email(rec["recipient_email"])
            status, error = send_outreach_email(recipient_email, subject, body, rec.get("recipient_name") or recipient_email)
            if status == "sent":
                sent += 1
                sent_at = now_iso()
            elif status.startswith("queued"):
                queued += 1
                sent_at = None
            else:
                failed += 1
                sent_at = None

            cur = conn.execute(
                """
                INSERT INTO email_outreach_logs(client_id, user_id, recipient_type, recipient_name, recipient_email, email_type,
                    subject, body, status, error, created_by, created_at, sent_at)
                VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)
                """,
                (rec["client_id"], rec["user_id"], rec["recipient_type"], rec["recipient_name"], recipient_email, email_type,
                 subject, body, status, error, user["username"], now_iso(), sent_at),
            )
            results.append({"id": cur.lastrowid, "email": recipient_email, "status": status, "error": error})
        conn.commit()

    log_action(user["username"], "send_customer_outreach", "email_outreach", None, f"type={email_type}; sent={sent}; queued={queued}; failed={failed}")
    if queued and not sent and not failed:
        message = f"Prepared {queued} email(s). Email provider is not configured, so they were logged but not sent."
    else:
        message = f"Outreach completed: sent {sent}, queued {queued}, failed {failed}."
    return {"message": message, "sent": sent, "queued": queued, "failed": failed, "results": results}


@app.post("/api/demo-data")
def demo_data(count: int = 50, user: Dict[str, Any] = Depends(require_roles("Admin"))) -> Dict[str, Any]:
    first = ["Aarav", "Nisha", "Karthik", "Meera", "Rohan", "Priya", "Vikram", "Ishita", "Arjun", "Sneha", "Rahul", "Ananya", "Dev", "Kavya", "Siddharth", "Pooja", "Manoj", "Ritika", "Aditya", "Divya"]
    last = ["Menon", "Rao", "Iyer", "Shah", "Das", "Nair", "Hegde", "Roy", "Kumar", "Bhat", "Shetty", "Gupta", "Patil", "Reddy", "Kulkarni"]
    roles = [
        ("React Developer", "React, TypeScript, API integration, responsive UI, design system"),
        ("Python Backend Developer", "Python, FastAPI, PostgreSQL, API security, deployment"),
        ("QA Automation Engineer", "Playwright, Selenium, API testing, test cases, UAT"),
        ("UI/UX Designer", "Figma, user journeys, wireframes, design system, mobile UX"),
        ("DevOps Engineer", "Docker, GitHub Actions, Railway, cloud deployment, CI/CD"),
        ("AI/ML Engineer", "LLM, RAG, Python, vector database, model evaluation"),
        ("Business Analyst", "Requirements, BRD, user stories, UAT, stakeholder management"),
        ("Data Engineer", "SQL, Python, ETL, data pipeline, Power BI"),
        ("Java Developer", "Java, Spring Boot, SQL, REST API, microservices"),
        ("Mobile Developer", "React Native, Flutter, API integration, mobile app deployment"),
    ]
    created = 0
    with get_db() as conn:
        for i in range(max(1, min(count, 250))):
            name = f"{first[i % len(first)]} {last[(i * 3) % len(last)]}"
            role_title, role_def = roles[i % len(roles)]
            skills = extract_skills(role_def)
            exp = 2.5 + (i % 9) + (0.5 if i % 2 else 0)
            risk_phrase = "" if i % 11 else " sample resume dummy candidate "
            resume_text = f"""{name}
Email: {name.lower().replace(' ', '.')}@example.com
Phone: +91 90000 {10000+i:05d}
Location: Bengaluru
Professional Summary: {exp:g} years of experience as {role_title}. Delivered production projects for Retail, SaaS, Government and Banking clients.
Skills: {role_def}, SQL, Agile, GitHub.
Projects: Built modules, APIs, dashboards, deployments, test automation and client-ready documentation from 2020 to 2025.
Certification: Internal 1Resource screening ready.{risk_phrase}
"""
            analysis = analyze_resume_text(resume_text, role_def, role_title)
            code = generate_candidate_code(conn)
            score = max(65, min(95, analysis["fit_score"] - (i % 7) + 5))
            status = status_from_score(score)
            cur = conn.execute(
                """
                INSERT INTO candidates(candidate_code, full_name, email, phone, location, current_status, availability_date, available_by_date, notice_period_days,
                    employment_type, source, recruiter_owner, total_experience, relevant_experience, primary_skill,
                    secondary_skills, domain_exposure, proficiency, certifications, portfolio_url, expected_rate, negotiated_rate,
                    internal_level, resume_text, status, ml_rating_score, ml_rating_level, fake_risk_score, fake_risk_level,
                    fake_risk_reasons, skill_matches, skill_gaps, last_role_definition, resume_count, created_date, created_at, updated_at)
                VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
                """,
                (code, name, f"{name.lower().replace(' ', '.')}@example.com", f"+91-90000-{10000+i:05d}", "Bengaluru",
                 ["Available", "Bench", "Notice Period", "Freelance"][i % 4], "Immediate" if i % 4 in [0,1] else "15 days", (datetime.utcnow() + timedelta(days=(0 if i % 4 in [0,1] else 15 + (i % 3) * 15))).date().isoformat(), 0 if i % 4 in [0,1] else 15 + (i % 3) * 15, "Contract", "Demo Data",
                 user["full_name"], exp, max(1, exp-1), skills[0] if skills else role_title, ", ".join(skills[1:]),
                 ", ".join(extract_domains(resume_text)), "Advanced" if exp >= 5 else "Intermediate", "", "", 80000 + i * 2500,
                 76000 + i * 2400, "L3 - Senior" if exp >= 5 else "L2 - Mid-level", build_standard_summary(name, skills, extract_domains(resume_text), exp, role_title, analysis["fit_score"], analysis["fake_risk_score"]),
                 status, analysis["fit_score"], analysis["rating_level"], analysis["fake_risk_score"], analysis["fake_risk_level"],
                 " | ".join(analysis["fake_risk_reasons"]), ", ".join(analysis["skill_matches"]), ", ".join(analysis["skill_gaps"]), role_def, 1, today_date(), now_iso(), now_iso())
            )
            candidate_id = int(cur.lastrowid)
            conn.execute(
                """
                INSERT INTO resumes(candidate_id, role_title, role_definition, resume_file_name, resume_file_path, resume_text, extracted_json,
                    fit_score, rating_level, fake_risk_score, fake_risk_level, fake_risk_reasons, skill_matches, skill_gaps, source, created_at)
                VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
                """,
                (candidate_id, role_title, role_def, f"demo_resume_{candidate_id}.txt", "", resume_text, json.dumps(analysis), analysis["fit_score"], analysis["rating_level"], analysis["fake_risk_score"], analysis["fake_risk_level"], " | ".join(analysis["fake_risk_reasons"]), ", ".join(analysis["skill_matches"]), ", ".join(analysis["skill_gaps"]), "Demo Data", now_iso())
            )
            conn.execute(
                """
                INSERT INTO assessments(candidate_id, technical_score, project_score, practical_score, communication_score,
                    client_readiness_score, cost_fitment_score, availability_score, total_score, evaluator_name, recommendation, remarks, created_at)
                VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?)
                """,
                (candidate_id, min(30, 18 + i % 12), min(15, 9 + i % 6), min(20, 12 + i % 8), min(10, 6 + i % 5),
                 min(10, 6 + i % 5), min(10, 6 + i % 5), 5 if i % 4 in [0,1] else 3, score, user["full_name"], status,
                 "Demo screening record generated for portal testing.", now_iso())
            )
            created += 1
        # Create demand-side sample requests so supply-demand analytics are visible immediately.
        demand_samples = [
            ("ADM", "Launchpad", "React Developer", "React, TypeScript, API integration, responsive dashboard UI", "React, TypeScript, JavaScript", "Data & Analytics", "Bengaluru", "Hot", 2, 275000, 180000),
            ("GovTech Client", "NFC Scheme Portal", "Python Backend Developer", "Python, FastAPI, PostgreSQL, secure APIs, file uploads", "Python, FastAPI, SQL", "Government", "Bengaluru", "High", 2, 300000, 190000),
            ("Retail SaaS", "QR Pedigree", "QA Automation Engineer", "Playwright, Selenium, API testing, regression suite", "QA Automation, Manual Testing", "Retail", "Remote", "Medium", 1, 180000, 120000),
            ("AI Client", "AI Governance", "AI/ML Engineer", "LLM, RAG, Python, evaluation, model routing", "AI/ML, Python", "SaaS", "Hybrid", "Hot", 1, 420000, 260000),
            ("Transport", "Biometric Verification", "DevOps Engineer", "Docker, CI/CD, Railway deployment, cloud monitoring", "Docker, DevOps, AWS", "Government", "Bengaluru", "High", 1, 240000, 160000),
            ("Enterprise", "Data Marketplace", "Business Analyst", "Requirements, user stories, UAT, stakeholder workshops", "Business Analyst, Project Management", "SaaS", "Hybrid", "Medium", 1, 220000, 140000),
            ("FoodFlow", "POS Platform", "Mobile Developer", "React Native, Flutter, API integration, mobile order flow", "Mobile, React, JavaScript", "Retail", "Remote", "Medium", 1, 250000, 160000),
            ("Analytics", "Executive Dashboard", "Data Engineer", "SQL, Python, ETL, Power BI dashboard data model", "Data Engineering, SQL, Power BI", "Banking", "Bengaluru", "High", 1, 320000, 210000),
        ]
        demand_created = 0
        for client, project, role_title, role_def, req_skills, domain, location, priority, positions, rate, cost in demand_samples:
            code = generate_demand_code(conn)
            conn.execute(
                """
                INSERT INTO demand_requests(demand_code, client_name, project_name, role_title, role_definition, required_skills,
                    domain, location, work_mode, priority, status, number_of_positions, target_customer_rate,
                    max_internal_cost, start_date, duration_weeks, created_by, created_date, created_at, updated_at)
                VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
                """,
                (code, client, project, role_title, role_def, req_skills, domain, location, "Hybrid" if location != "Remote" else "Remote",
                 priority, "Open", positions, rate, cost, "Immediate", 12, user["username"], today_date(), now_iso(), now_iso()),
            )
            demand_created += 1
        conn.commit()
    audit_status = "logged"
    try:
        log_action(user["username"], "create_50_demo_data", "candidate", None, f"created={created}; demand_created={demand_created}")
    except Exception as exc:
        # Demo data creation has already been committed. Do not convert a successful seed into a failed API response
        # only because the non-critical audit insert failed on a specific database runtime.
        audit_status = f"audit_log_failed: {str(exc)[:180]}"
    return {
        "ok": True,
        "created": created,
        "demand_created": demand_created,
        "audit_status": audit_status,
        "message": f"Created {created} test candidate records and {demand_created} demand requests"
    }


# Serve React app if frontend has been built.
FRONTEND_DIST_OPTIONS = [
    os.path.abspath(os.path.join(BASE_DIR, "..", "frontend", "dist")),
    os.path.abspath(os.path.join(BASE_DIR, "frontend_dist")),
]
FRONTEND_DIST = next((p for p in FRONTEND_DIST_OPTIONS if os.path.exists(p)), None)
if FRONTEND_DIST:
    app.mount("/assets", StaticFiles(directory=os.path.join(FRONTEND_DIST, "assets")), name="assets")

    @app.get("/{full_path:path}")
    def spa(full_path: str):
        index_path = os.path.join(FRONTEND_DIST, "index.html")
        if full_path and os.path.exists(os.path.join(FRONTEND_DIST, full_path)):
            return FileResponse(os.path.join(FRONTEND_DIST, full_path))
        return FileResponse(index_path)
